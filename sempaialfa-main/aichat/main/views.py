from django.shortcuts import render
from django.http import JsonResponse, HttpResponse, Http404
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
import time
import json
import re
import requests
import base64
import logging
import threading
import csv
from datetime import datetime, timedelta
from django.conf import settings
from django.utils import timezone
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
import re
from .file_processor import process_file
from .models import ChatRequest, ChatHistory, Metric, UserActivity
from .content_moderator import ContentModerator
from .metrics_calculator import MetricsCalculator
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login as django_login

logger = logging.getLogger(__name__)


def index(request):
    """Главная страница"""
    # Добавляем timestamp, чтобы браузер не кешировал старую версию
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/index.html', context)


def cabinet(request):
    """Страница личного кабинета"""
    try:
        # Добавляем timestamp, чтобы браузер не кешировал старую версию
        context = {
            'cache_version': int(time.time())
        }
        response = render(request, 'main/cabinet.html', context)
        # Проверяем, что ответ не пустой
        if not response.content:
            logger.error("Пустой ответ при рендеринге cabinet.html")
            return HttpResponse("Ошибка: пустой ответ сервера", status=500)
        return response
    except Exception as e:
        logger.error(f"Ошибка при загрузке страницы cabinet: {str(e)}", exc_info=True)
        return HttpResponse(f"Ошибка загрузки страницы: {str(e)}", status=500)


def chat(request):
    """Страница чата"""
    # Добавляем timestamp, чтобы браузер не кешировал старую версию
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/chat.html', context)


def transfer(request):
    """Страница перевода средств"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/transfer.html', context)


def receipts(request):
    """Страница чеков"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/receipts.html', context)


def utilities(request):
    """Страница коммунальных услуг"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/utilities.html', context)


def taxes(request):
    """Страница налоговых выплат"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/taxes.html', context)


def calendar(request):
    """Страница календаря"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/calendar.html', context)


def documents(request):
    """Страница документов"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/documents.html', context)


def inventory(request):
    """Страница инвентаризации"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/inventory.html', context)


def employees(request):
    """Страница сотрудников"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/employees.html', context)


def support(request):
    """Страница поддержки"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/support.html', context)


def mail(request):
    """Страница почты"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/mail.html', context)


def login(request):
    """Страница входа/регистрации"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/login.html', context)


@csrf_exempt
@require_http_methods(["POST"])
def login_api(request):
    """API для авторизации пользователей"""
    
    try:
        data = json.loads(request.body)
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        
        if not email or not password:
            return JsonResponse({'success': False, 'error': 'Email и пароль обязательны'})
        
        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            return JsonResponse({
                'success': False, 
                'error': 'Неверный email или пароль',
                'use_localstorage': True
            })
        
        user = authenticate(request, username=user.username, password=password)
        
        if user is not None:
            django_login(request, user)
            is_admin = user.is_staff or user.is_superuser
            
            return JsonResponse({
                'success': True,
                'is_admin': is_admin,
                'username': user.username,
                'email': user.email
            })
        else:
            return JsonResponse({'success': False, 'error': 'Неверный пароль'})
            
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Неверный формат данных'})
    except Exception as e:
        logger.error(f'Ошибка в login_api: {str(e)}')
        return JsonResponse({'success': False, 'error': f'Ошибка сервера: {str(e)}'})


def feedback(request):
    """Страница обратной связи"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/feedback.html', context)


def scenarios(request):
    """Страница сценариев использования"""
    context = {
        'cache_version': int(time.time())
    }
    return render(request, 'main/scenarios.html', context)


@csrf_exempt
@require_http_methods(["GET", "POST"])
def check_lm_studio_connection(request):
    """API endpoint для проверки подключения к OpenRouter"""
    try:
        OPENROUTER_API_KEY = getattr(settings, 'OPENROUTER_API_KEY', '')
        OPENROUTER_URL = getattr(settings, 'OPENROUTER_URL', 'https://openrouter.ai/api/v1/chat/completions')
        OPENROUTER_MODEL = getattr(settings, 'OPENROUTER_MODEL', 'deepseek/deepseek-r1')
        
        if not OPENROUTER_API_KEY:
            return JsonResponse({
                'success': False,
                'server_available': False,
                'api_available': False,
                'model_available': False,
                'url': OPENROUTER_URL,
                'error': 'API ключ OpenRouter не настроен. Установите переменную окружения OPENROUTER_API_KEY.',
                'recommendations': [
                    'Проверьте, что в файле .env установлен OPENROUTER_API_KEY',
                    'Получите API ключ на https://openrouter.ai/keys',
                    'Убедитесь, что переменная окружения правильно экспортирована'
                ]
            })
        
        # Проверяем API, отправляя тестовый запрос
        api_available = False
        model_available = False
        error_details = None
        
        try:
            headers = {
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
                "HTTP-Referer": request.build_absolute_uri('/') if hasattr(request, 'build_absolute_uri') else "https://localhost",
            }
            
            test_payload = {
                "model": OPENROUTER_MODEL,
                "messages": [{"role": "user", "content": "test"}],
                "max_tokens": 1
            }
            
            api_response = requests.post(OPENROUTER_URL, headers=headers, json=test_payload, timeout=10)
            
            if api_response.status_code == 200:
                api_available = True
                model_available = True
            elif api_response.status_code == 401:
                api_available = True
                model_available = False
                error_details = "Неверный API ключ. Проверьте правильность OPENROUTER_API_KEY."
            elif api_response.status_code == 404:
                api_available = True
                model_available = False
                error_details = f"Модель '{OPENROUTER_MODEL}' не найдена. Проверьте название модели на https://openrouter.ai/models"
            else:
                api_available = True
                try:
                    error_data = api_response.json()
                    error_message = error_data.get('error', {}).get('message', '') if isinstance(error_data.get('error'), dict) else str(error_data.get('error', ''))
                    if error_message:
                        error_details = f"Статус {api_response.status_code}: {error_message}"
                    else:
                        error_details = f"Статус {api_response.status_code}: {error_data}"
                except:
                    error_details = f"Статус {api_response.status_code}: {api_response.text[:200]}"
        except requests.exceptions.ConnectionError:
            api_available = False
            error_details = "Не удалось подключиться к OpenRouter API. Проверьте интернет-соединение."
        except requests.exceptions.Timeout:
            api_available = False
            error_details = "Превышено время ожидания ответа от OpenRouter"
        except Exception as e:
            api_available = False
            error_details = str(e)
        
        return JsonResponse({
            'success': api_available and model_available,
            'server_available': True,  # OpenRouter всегда доступен (облачный сервис)
            'api_available': api_available,
            'model_available': model_available,
            'url': OPENROUTER_URL,
            'error': error_details,
            'recommendations': [] if (api_available and model_available) else [
                'Проверьте, что API ключ OpenRouter установлен в переменной окружения OPENROUTER_API_KEY',
                'Получите API ключ на https://openrouter.ai/keys',
                f'Проверьте, что модель {OPENROUTER_MODEL} доступна на https://openrouter.ai/models',
                'Убедитесь, что у вас есть активное интернет-соединение',
                'Проверьте, что на вашем аккаунте OpenRouter достаточно баланса'
            ]
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Ошибка проверки: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def register_user(request):
    """API endpoint для регистрации пользователей"""
    try:
        data = json.loads(request.body)
        organization = data.get('organization', '').strip()
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        confirm_password = data.get('confirmPassword', '')
        
        # Проверяем, что все обязательные поля заполнены
        if not email or not password:
            return JsonResponse({
                'success': False,
                'error': 'Email и пароль обязательны'
            }, status=400)
        
        if password != confirm_password:
            return JsonResponse({
                'success': False,
                'error': 'Пароли не совпадают'
            }, status=400)
        
        if len(password) < 6:
            return JsonResponse({
                'success': False,
                'error': 'Пароль должен содержать минимум 6 символов'
            }, status=400)
        
        # Проверяем, нет ли уже пользователя с таким email
        if User.objects.filter(email=email).exists():
            return JsonResponse({
                'success': False,
                'error': 'Пользователь с таким email уже зарегистрирован'
            }, status=400)
        
        # Создаем пользователя
        # Используем часть email до @ как username
        username = email.split('@')[0]
        # Если username занят, добавляем цифру в конце
        base_username = username
        counter = 1
        while User.objects.filter(username=username).exists():
            username = f"{base_username}{counter}"
            counter += 1
        
        user = User.objects.create_user(
            username=username,
            email=email,
            password=password,
            first_name=organization  # Сохраняем название организации в first_name
        )
        
        logger.info(f"Создан новый пользователь: {email} (username: {username})")
        
        # Создаем записи для метрик и аналитики
        try:
            # Создаем запись в истории чатов для отслеживания регистрации
            registration_chat_id = f"registration_{int(time.time() * 1000)}"
            ChatHistory.objects.create(
                user_email=email,
                chat_id=registration_chat_id,
                title='Регистрация пользователя',
                messages=[{
                    'text': f'Регистрация пользователя: {organization}',
                    'isUser': False,
                    'timestamp': timezone.now().isoformat(),
                    'type': 'registration'
                }],
                ai_actions=[],
                total_messages=1,
                total_user_messages=0,
                total_ai_messages=1,
                total_actions=0,
                last_message_at=timezone.now()
            )
            logger.info(f"Создана запись ChatHistory для регистрации пользователя {email}")
            
            # Создаем запись активности пользователя для детального отслеживания
            ip_address = request.META.get('REMOTE_ADDR', '')
            user_agent = request.META.get('HTTP_USER_AGENT', '')
            UserActivity.objects.create(
                user=user,
                user_email=email,
                activity_type='registration',
                activity_data={
                    'organization': organization,
                    'username': username
                },
                ip_address=ip_address if ip_address else None,
                user_agent=user_agent
            )
            logger.info(f"Создана запись UserActivity для регистрации пользователя {email}")
        except Exception as e:
            logger.error(f"Ошибка при создании записей для регистрации: {str(e)}", exc_info=True)
            # Если метрики не записались, это не должно прерывать регистрацию
        
        return JsonResponse({
            'success': True,
            'message': 'Пользователь успешно зарегистрирован',
            'user': {
                'id': user.id,
                'email': user.email,
                'username': user.username,
                'organization': user.first_name
            }
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Неверный формат данных'
        }, status=400)
    except Exception as e:
        logger.error(f"Ошибка при регистрации: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': 'Внутренняя ошибка сервера'
        }, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def get_user_data(request):
    """API endpoint для получения данных пользователя"""
    try:
        data = json.loads(request.body)
        email = data.get('email', '')
        
        if not email:
            return JsonResponse({
                'success': False,
                'error': 'Email не указан'
            }, status=400)
        
        # Данные хранятся на фронтенде в localStorage
        return JsonResponse({
            'success': True,
            'message': 'Данные должны быть переданы из фронтенда'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Ошибка: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def create_calendar_event(request):
    """API endpoint для создания события в календаре"""
    try:
        data = json.loads(request.body)
        email = data.get('email', '')
        title = data.get('title', '')
        date = data.get('date', '')
        description = data.get('description', '')
        
        if not all([email, title, date]):
            return JsonResponse({
                'success': False,
                'error': 'Не все обязательные поля заполнены'
            }, status=400)
        
        # Формируем объект события, фронтенд сохранит его в localStorage
        event = {
            'id': str(int(time.time() * 1000)),
            'title': title,
            'date': date,
            'description': description,
            'notified': False
        }
        
        return JsonResponse({
            'success': True,
            'event': event
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Ошибка: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def manage_calendar_event(request):
    """Универсальный API endpoint для управления событиями календаря (создание, обновление, удаление)"""
    try:
        data = json.loads(request.body)
        action = data.get('action')  # Может быть CREATE_EVENT, UPDATE_EVENT или DELETE_EVENT
        email = data.get('email', '')
        event_id = data.get('event_id', '')
        title = data.get('title', '')
        date = data.get('date', '')
        description = data.get('description', '')
        
        if not email:
            return JsonResponse({
                'success': False,
                'error': 'Email не указан'
            }, status=400)
        
        # Формируем структурированный ответ для фронтенда
        result = {
            'success': True,
            'action': action,
            'email': email
        }
        
        if action == 'create':
            if not title or not date:
                return JsonResponse({
                    'success': False,
                    'error': 'Для создания события необходимы title и date'
                }, status=400)
            result['event'] = {
                'id': str(int(time.time() * 1000)),
                'title': title,
                'date': date,
                'description': description or '',
                'notified': False
            }
        elif action == 'update':
            if not event_id:
                return JsonResponse({
                    'success': False,
                    'error': 'Для обновления события необходим event_id'
                }, status=400)
            result['event'] = {
                'id': event_id,
                'title': title,
                'date': date,
                'description': description
            }
        elif action == 'delete':
            if not event_id:
                return JsonResponse({
                    'success': False,
                    'error': 'Для удаления события необходим event_id'
                }, status=400)
            result['event_id'] = event_id
        else:
            return JsonResponse({
                'success': False,
                'error': f'Неизвестное действие: {action}'
            }, status=400)
        
        return JsonResponse(result)
    except Exception as e:
        logger.error(f"Ошибка в manage_calendar_event: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка: {str(e)}'
        }, status=500)


def format_user_context(user_data):
    """Форматирует данные пользователя для передачи в AI"""
    context_parts = []
    
    # Балансы счетов
    balance1 = user_data.get('accountBalance', 0)
    balance2 = user_data.get('accountBalance2', 0)
    if balance1 or balance2:
        context_parts.append(f"Балансы счетов: Счет 1 - {balance1:,.0f} ₽, Счет 2 - {balance2:,.0f} ₽")
    
    # Чеки (для визуализации)
    receipts = user_data.get('receipts', [])
    if receipts:
        total_receipts = len(receipts)
        total_amount = sum(r.get('amount', 0) for r in receipts if isinstance(r.get('amount'), (int, float)))
        context_parts.append(f"\nЧЕКИ: {total_receipts} операций, сумма {total_amount:,.0f} ₽")
        
        # Группируем операции по типам и датам для графиков
        operations_by_type = {}
        operations_by_date = {}
        for r in receipts:
            op_type = r.get('operationType', 'Операция')
            amount = r.get('amount', 0)
            date = r.get('date', '')
            
            # Группировка по типам
            if op_type not in operations_by_type:
                operations_by_type[op_type] = 0
            operations_by_type[op_type] += amount
            
            # Группировка по датам, если указаны
            if date:
                date_key = date.split('T')[0] if 'T' in date else date.split(' ')[0]
                if date_key not in operations_by_date:
                    operations_by_date[date_key] = 0
                operations_by_date[date_key] += amount
        
        # Добавляем данные для графиков
        if operations_by_type:
            context_parts.append("Распределение по типам операций:")
            for op_type, total in operations_by_type.items():
                context_parts.append(f"  - {op_type}: {total:,.0f} ₽")
        
        # Последние операции
        if receipts:
            recent = receipts[-5:]  # Берем последние 5
            context_parts.append("Последние операции:")
            for r in recent:
                op_type = r.get('operationType', 'Операция')
                amount = r.get('amount', 0)
                date = r.get('date', 'Не указана')
                context_parts.append(f"  - {op_type}: {amount:,.0f} ₽ ({date})")
    
    # Инвентаризация (для визуализации)
    inventory = user_data.get('inventory', [])
    inventory_folders = user_data.get('inventoryFolders', [])
    
    if inventory:
        context_parts.append(f"\nИНВЕНТАРИЗАЦИЯ: {len(inventory)} позиций")
        
        # Создаем соответствие между id папок и их названиями
        folder_map = {}
        if inventory_folders:
            for folder in inventory_folders:
                folder_id = folder.get('id') or folder.get('folderId')
                folder_name = folder.get('name', 'Без названия')
                if folder_id:
                    folder_map[str(folder_id)] = folder_name
        
        # Группируем товары по папкам (категориям)
        categories = {}
        categories_value = {}
        categories_items = {}  # Храним товары по категориям
        
        for item in inventory:
            # Получаем название папки для товара
            folder_id = str(item.get('folderId') or item.get('folder') or '')
            if folder_id and folder_id in folder_map:
                cat = folder_map[folder_id]
            elif item.get('folderName'):
                cat = item.get('folderName')
            elif item.get('folder'):
                cat = item.get('folder')
            else:
                cat = 'Без категории'
            
            quantity = item.get('quantity', 0)
            price = item.get('price', 0)
            value = quantity * price if isinstance(quantity, (int, float)) and isinstance(price, (int, float)) else 0
            item_name = item.get('name', 'Без названия')
            
            if cat not in categories:
                categories[cat] = 0
                categories_value[cat] = 0
                categories_items[cat] = []
            
            categories[cat] += quantity
            categories_value[cat] += value
            categories_items[cat].append({
                'name': item_name,
                'quantity': quantity,
                'price': price,
                'value': value
            })
        
        # Структурированный вывод: каждая папка с товарами внутри
        if categories:
            context_parts.append("СТРУКТУРА ИНВЕНТАРИЗАЦИИ (папки и товары):")
            for cat, count in categories.items():
                value = categories_value.get(cat, 0)
                items_in_cat = categories_items.get(cat, [])
                context_parts.append(f"ПАПКА: '{cat}' - {count} шт., общая стоимость {value:,.2f} ₽")
                # Перечисляем все товары в папке
                for item in items_in_cat:
                    context_parts.append(f"  - ТОВАР: '{item['name']}' | Количество: {item['quantity']} | Цена: {item['price']:,.2f} ₽ | Стоимость: {item['value']:,.2f} ₽")
        
        # Добавляем общее распределение по категориям
        if categories:
            context_parts.append("\nРаспределение по категориям (кратко):")
            for cat, count in categories.items():
                value = categories_value.get(cat, 0)
                context_parts.append(f"  - {cat}: {count} шт., стоимость {value:,.2f} ₽")
        
        # Общая стоимость инвентаризации
        all_items_value = sum(
            item.get('quantity', 0) * item.get('price', 0) 
            for item in inventory 
            if isinstance(item.get('quantity'), (int, float)) and isinstance(item.get('price'), (int, float))
        )
        if all_items_value > 0:
            context_parts.append(f"Общая стоимость инвентаризации: {all_items_value:,.2f} ₽")
        
        # Топ товары по стоимости
        items_with_value = [
            {
                'name': item.get('name', 'Без названия'),
                'category': item.get('folder') or item.get('folderId') or 'Без категории',
                'quantity': item.get('quantity', 0),
                'price': item.get('price', 0),
                'value': (item.get('quantity', 0) * item.get('price', 0)) if isinstance(item.get('quantity'), (int, float)) and isinstance(item.get('price'), (int, float)) else 0
            }
            for item in inventory
        ]
        items_with_value.sort(key=lambda x: x['value'], reverse=True)
        if items_with_value:
            context_parts.append("Топ-5 товаров по стоимости:")
            for item in items_with_value[:5]:
                context_parts.append(f"  - {item['name']} ({item['category']}): {item['value']:,.2f} ₽")
        
        # Полный список товаров с названиями для AI
        # Важно: AI должен использовать названия, а не ID
        if inventory:
            context_parts.append("\nПОЛНЫЙ СПИСОК ТОВАРОВ (ВАЖНО: используй НАЗВАНИЯ, а не ID):")
            for item in inventory[:30]:  # Берем 30 для контекста
                item_name = item.get('name', 'Без названия')
                # Получаем название папки
                folder_id = str(item.get('folderId') or item.get('folder') or '')
                if folder_id and folder_id in folder_map:
                    item_cat = folder_map[folder_id]
                elif item.get('folderName'):
                    item_cat = item.get('folderName')
                elif item.get('folder'):
                    item_cat = item.get('folder')
                else:
                    item_cat = 'Без категории'
                
                item_qty = item.get('quantity', 0)
                item_price = item.get('price', 0)
                item_value = item_qty * item_price if isinstance(item_qty, (int, float)) and isinstance(item_price, (int, float)) else 0
                # Явно указываем названия, а не ID
                context_parts.append(f"  - ТОВАР: '{item_name}' | ПАПКА: '{item_cat}' | Количество: {item_qty} | Цена: {item_price:,.2f} ₽ | Стоимость: {item_value:,.2f} ₽")
            if len(inventory) > 30:
                context_parts.append(f"  ... и еще {len(inventory) - 30} товаров")
            context_parts.append("\nКРИТИЧЕСКИ ВАЖНО:")
            context_parts.append("- Всегда используй НАЗВАНИЯ товаров из поля 'ТОВАР', НИКОГДА не используй числовые ID!")
            context_parts.append("- Всегда используй НАЗВАНИЯ папок из поля 'ПАПКА', НИКОГДА не используй числовые ID папок!")
            context_parts.append("- Когда пользователь просит показать инвентаризацию, перечисляй каждую ПАПКУ с товарами внутри!")
    
    # Сотрудники (для визуализации)
    employees = user_data.get('employees', [])
    if employees:
        total_salary = sum(e.get('salary', 0) for e in employees if isinstance(e.get('salary'), (int, float)))
        avg_salary = total_salary / len(employees) if employees else 0
        context_parts.append(f"\nСОТРУДНИКИ: {len(employees)} человек, фонд {total_salary:,.0f} ₽, средняя зарплата {avg_salary:,.0f} ₽")
        
        # Полный список сотрудников для графиков
        context_parts.append("Список сотрудников с зарплатами:")
        for emp in employees:
            fio = emp.get('fio', 'Не указано')
            salary = emp.get('salary', 0)
            position = emp.get('position', 'Не указана')
            context_parts.append(f"  - {fio} ({position}): {salary:,.0f} ₽")
    
    # Календарь
    calendar_events = user_data.get('calendarEvents', [])
    if calendar_events:
        context_parts.append(f"\nКАЛЕНДАРЬ: {len(calendar_events)} событий")
        upcoming = [e for e in calendar_events if e.get('date')]
        upcoming.sort(key=lambda x: x.get('date', ''))
        upcoming = upcoming[:3]  # Берем только 3 ближайших
        if upcoming:
            for e in upcoming:
                event_id = e.get('id', '')
                title = e.get('title', 'Событие')
                date = e.get('date', '')
                # Извлекаем дату в формате ISO
                iso_date = ''
                if date:
                    try:
                        event_date = datetime.fromisoformat(date.replace('Z', '+00:00'))
                        iso_date = event_date.strftime('%Y-%m-%dT%H:%M')
                    except:
                        if 'T' in date:
                            iso_date = date[:16]
                        else:
                            iso_date = date
                context_parts.append(f"  - ID: {event_id} | '{title}' | {iso_date}")
    
    # Налоги
    taxes_data = user_data.get('taxesData', {})
    if taxes_data:
        context_parts.append(f"\nНАЛОГИ:")
        tax_names = {
            'profit': 'Налог на прибыль',
            'vat': 'НДС',
            'property': 'Налог на имущество',
            'insurance': 'Страховые взносы'
        }
        total_tax_debt = 0
        for tax_id, tax_data in taxes_data.items():
            debt = tax_data.get('debt', 0)
            if isinstance(debt, (int, float)) and debt > 0:
                tax_name = tax_names.get(tax_id, tax_id)
                context_parts.append(f"  - {tax_name}: {debt:,.0f} ₽")
                total_tax_debt += debt
        if total_tax_debt > 0:
            context_parts.append(f"Общая задолженность по налогам: {total_tax_debt:,.0f} ₽")
    
    # Коммунальные услуги
    utilities_data = user_data.get('utilitiesData', {})
    if utilities_data:
        context_parts.append(f"\nКОММУНАЛЬНЫЕ УСЛУГИ:")
        util_names = {
            'electricity': 'Электричество',
            'water': 'Водоснабжение',
            'heating': 'Отопление',
            'waste': 'Вывоз ТКО',
            'security': 'Охранные услуги',
            'internet': 'Интернет'
        }
        total_utilities_debt = 0
        for util_id, util_data in utilities_data.items():
            debt = util_data.get('debt', 0)
            if isinstance(debt, (int, float)) and debt > 0:
                util_name = util_names.get(util_id, util_id)
                context_parts.append(f"  - {util_name}: {debt:,.0f} ₽")
                total_utilities_debt += debt
        if total_utilities_debt > 0:
            context_parts.append(f"Общая задолженность по коммунальным услугам: {total_utilities_debt:,.0f} ₽")
    
    # Документы
    documents = user_data.get('documents', [])
    if documents:
        context_parts.append(f"\nДОКУМЕНТЫ: {len(documents)} файлов")
        for doc in documents[:3]:  # Берем первые 3, чтобы не перегружать контекст
            name = doc.get('name', 'Без названия')
            size = doc.get('size', 0)
            doc_type = doc.get('type', 'неизвестный тип')
            context_parts.append(f"  - {name} ({doc_type}, {size} байт)")
            # Содержимое документов передается отдельно через файлы
    
    return "\n".join(context_parts) if context_parts else "Данные отсутствуют"


def find_event_smart(calendar_events, identifier):
    """
    Умный поиск события в календаре по различным критериям:
    - Точный ID
    - Точное название
    - Часть названия
    - Описание события
    - Дата события
    """
    if not calendar_events or not identifier:
        return None
    
    identifier = identifier.strip()
    identifier_lower = identifier.lower().strip()
    
    # Сначала ищем по точному ID
    for event in calendar_events:
        if str(event.get('id', '')) == identifier:
            logger.info(f"Событие найдено по точному ID: {identifier}")
            return str(event.get('id', ''))
    
    # Потом по точному названию (регистр не важен)
    for event in calendar_events:
        event_title = event.get('title', '').lower().strip()
        if event_title == identifier_lower:
            logger.info(f"Событие найдено по точному названию: '{event.get('title')}'")
            return str(event.get('id', ''))
    
    # Поиск по части названия
    for event in calendar_events:
        event_title = event.get('title', '').lower().strip()
        if identifier_lower in event_title or event_title in identifier_lower:
            logger.info(f"Событие найдено по части названия: '{event.get('title')}'")
            return str(event.get('id', ''))
    
    # Поиск по описанию
    for event in calendar_events:
        event_description = event.get('description', '').lower().strip()
        if event_description and identifier_lower in event_description:
            logger.info(f"Событие найдено по описанию: '{event.get('title')}'")
            return str(event.get('id', ''))
    
    # Поиск по ключевым словам в названии
    identifier_words = [w for w in identifier_lower.split() if len(w) > 2]  # Короткие слова игнорируем
    if identifier_words:
        best_match = None
        best_score = 0
        
        for event in calendar_events:
            event_title = event.get('title', '').lower()
            event_desc = event.get('description', '').lower()
            event_text = f"{event_title} {event_desc}"
            
            score = 0
            for word in identifier_words:
                if word in event_text:
                    score += 1
            
            if score > best_score and score >= len(identifier_words) * 0.5:  # Нужно совпадение хотя бы половины слов
                best_score = score
                best_match = event
        
        if best_match:
            logger.info(f"Событие найдено по ключевым словам: '{best_match.get('title')}' (score: {best_score})")
            return str(best_match.get('id', ''))
    
    logger.warning(f"Событие не найдено по идентификатору: '{identifier}'")
    return None


def process_chat_request_async(request_id):
    """Асинхронная обработка запроса к AI в фоновом потоке"""
    try:
        from django.utils import timezone
        
        # Отслеживаем время начала обработки
        processing_start_time = timezone.now()
        llm_start_time = None
        
        chat_request = ChatRequest.objects.get(id=request_id)
        chat_request.status = ChatRequest.STATUS_PROCESSING
        chat_request.save()
        
        # Импортируем логику обработки из chat_api
        message = chat_request.message
        chat_history = chat_request.chat_history
        user_data = chat_request.user_data
        files = chat_request.files_data
        
        # Проверяем сообщение на модерацию
        message_blocked = False
        moderation_result = ContentModerator.check_message(message)
        if not moderation_result['allowed']:
            message_blocked = True
        
        # Настройки OpenRouter
        OPENROUTER_API_KEY = getattr(settings, 'OPENROUTER_API_KEY', '')
        OPENROUTER_URL = getattr(settings, 'OPENROUTER_URL', 'https://openrouter.ai/api/v1/chat/completions')
        OPENROUTER_MODEL = getattr(settings, 'OPENROUTER_MODEL', 'deepseek/deepseek-r1')
        
        if not OPENROUTER_API_KEY:
            chat_request.status = ChatRequest.STATUS_FAILED
            chat_request.error = 'API ключ OpenRouter не настроен. Установите переменную окружения OPENROUTER_API_KEY.'
            chat_request.save()
            return
        
        # Формируем историю сообщений для модели
        messages = []
        
        # Формируем контекст пользователя
        user_context = format_user_context(user_data)
        
        # Получаем текущую дату для вычисления относительных дат
        now = timezone.now()
        current_date_str = now.strftime('%Y-%m-%d')
        current_time_str = now.strftime('%H:%M')
        current_datetime_str = now.strftime('%Y-%m-%d %H:%M')
        current_date_readable = now.strftime('%d %B %Y года')
        
        # Добавляем системный промпт (сокращенная версия для скорости)
        system_prompt = f"""Ты - профессиональный бизнес-ассистент. Работай СТРОГО в деловом стиле. Отвечай кратко и точно.

ТЕКУЩАЯ ДАТА И ВРЕМЯ:
СЕЙЧАС: {current_datetime_str} ({current_date_readable})
Используй ЭТУ дату как базовую для вычисления относительных дат ("завтра", "через неделю" и т.д.).

ПОНИМАНИЕ ЕСТЕСТВЕННОГО ЯЗЫКА:
Пользователь пишет ОБЫЧНЫМ языком (например: "добавь встречу", "удали совещание", "перенеси на завтра").
ТЫ ДОЛЖЕН сам переводить эти просьбы в технические команды.
НИКОГДА НЕ ПРОСИ пользователя писать команды вручную (типа CREATE_EVENT или UPDATE_EVENT).
Просто выполняй действие и подтверждай его словами.

ВОЗМОЖНОСТИ И КОМАНДЫ (ДЛЯ ТЕБЯ):
- Анализ: чеки, инвентаризация, сотрудники, календарь, налоги, коммунальные услуги, документы, балансы
- Действия (пиши их на отдельной строке):
  * CREATE_EVENT: название|дата ISO|описание
  * DELETE_EVENT: название
  * UPDATE_EVENT: {{"action":"UPDATE_EVENT","event":"ID","title":"...","date":"YYYY-MM-DDTHH:mm","description":"..."}}
  * DELETE_DOCUMENT: название
  * RENAME_DOCUMENT: старое|новое
  * SEND_SUPPORT_MESSAGE: тема|текст

ВАЖНО ДЛЯ КАЛЕНДАРЯ:
1. Если пользователь пишет "запланируй", "добавь", "напомни" -> используй CREATE_EVENT.
   - Пример: "Напомни позвонить маме завтра в 5 вечера" -> вычисли дату завтра от {current_date_str} и создай CREATE_EVENT: Звонок маме|YYYY-MM-DDTHH:mm|Позвонить маме
2. Если пользователь пишет "удали", "отмени" -> используй DELETE_EVENT.
   - Пример: "Удали встречу с клиентом" -> DELETE_EVENT: Встреча с клиентом
3. Если пользователь пишет "перенеси", "измени" -> используй UPDATE_EVENT.
   - Пример: "Перенеси совещание на послезавтра на 10 утра" -> вычисли дату послезавтра от {current_date_str} и создай {{"action":"UPDATE_EVENT", "event":"Совещание", "date":"YYYY-MM-DDTHH:mm"}}

ОБЯЗАТЕЛЬНО вычисляй точную дату из относительных формулировок ОТНОСИТЕЛЬНО ТЕКУЩЕЙ ДАТЫ ({current_date_str}):
- "через неделю" = {current_date_str} + 7 дней
- "завтра" = {current_date_str} + 1 день
- "послезавтра" = {current_date_str} + 2 дня
- "через месяц" = {current_date_str} + 30 дней
- "в 15:00" = 15:00
- "в полдень" = 12:00
- "в 5 вечера" = 17:00
- "в 10 утра" = 10:00

Формат даты СТРОГО: YYYY-MM-DDTHH:mm

ВАЖНО:
- Команду пиши на отдельной строке.
- Пользователю отвечай вежливо: "Хорошо, я запланировал...", "Событие удалено.", "Встреча перенесена."
- НЕ показывай пользователю технические детали команд, если это не требуется для отладки.

ВИЗУАЛИЗАЦИЯ:
- Таблицы Markdown: | Колонка1 | Колонка2 |\n|:---|:---:|\n| Данные1 | Данные2 |
- Графики и диаграммы: используй простые команды [CHART_ТИП_ДАННЫХ:тип_графика]

ВИЗУАЛИЗАЦИЯ ДАННЫХ:
Когда пользователь просит показать график, диаграмму или визуализацию данных, используй ПРОСТЫЕ КОМАНДЫ:
[CHART_ТИП_ДАННЫХ:тип_графика]

Типы данных:
- RECEIPTS или ЧЕКИ - для графиков по чекам/операциям
- INVENTORY или ИНВЕНТАРИЗАЦИЯ - для графиков по инвентаризации
- EMPLOYEES или СОТРУДНИКИ - для графиков по сотрудникам
- TAXES или НАЛОГИ - для графиков по налогам
- UTILITIES или КОММУНАЛЬНЫЕ - для графиков по коммунальным услугам
- BALANCE или БАЛАНС - для графиков балансов счетов

Типы графиков:
- line: линейный график (для динамики по времени)
- bar: столбчатая диаграмма (для сравнения)
- pie: круговая диаграмма (для распределения)
- doughnut: кольцевая диаграмма (для распределения)
- horizontal: горизонтальная столбчатая (для топ-списков)

ПРИМЕРЫ ИСПОЛЬЗОВАНИЯ:
- "Покажи график расходов" → [CHART_RECEIPTS:pie]
- "Создай круговую диаграмму инвентаризации" → [CHART_INVENTORY:pie]
- "Сравни зарплаты сотрудников" → [CHART_EMPLOYEES:bar]
- "Покажи распределение задолженностей по налогам" → [CHART_TAXES:doughnut]
- "График коммунальных услуг" → [CHART_UTILITIES:pie]
- "Сравни балансы счетов" → [CHART_BALANCE:bar]

ВАЖНО: 
- Используй ТОЛЬКО простые команды [CHART_ТИП:тип_графика]
- НЕ создавай JSON вручную - система сама извлечет данные из контекста пользователя
- Если данных недостаточно, сообщи об этом пользователю
- Выбирай подходящий тип графика: pie/doughnut для распределения, bar для сравнения, line для динамики

СТИЛЬ: Деловой, формальный, без анекдотов и шуток. Для данных используй таблицы, для визуализации - графики.

Данные пользователя:
""" + user_context[:1000]  # Ограничиваем контекст до 1000 символов для оптимизации
        
        messages.append({
            "role": "system",
            "content": system_prompt
        })
        
        # Обрабатываем файлы, если они есть
        image_files = []  # Список изображений для отправки в vision модель
        file_contents = []  # Список текстового содержимого файлов
        
        if files:
            for i, file in enumerate(files):
                try:
                    file_name = file.get('name', f'Файл {i+1}')
                    file_type = file.get('type', 'неизвестный тип')
                    file_data = file.get('data', '')
                    
                    if not file_data:
                        file_contents.append(f"Файл {i+1} ({file_name}): [Файл пуст или данные не получены]")
                        continue
                    
                    # Обрабатываем файл с помощью модуля file_processor
                    try:
                        extracted_text, image_base64 = process_file(file_name, file_type, file_data)
                    except Exception as e:
                        logger.error(f"Ошибка при обработке файла '{file_name}': {str(e)}", exc_info=True)
                        error_msg = f"[Ошибка при обработке файла '{file_name}': {str(e)}]"
                        file_contents.append(f"Файл {i+1} ({file_name}): {error_msg}")
                        continue
                    
                    if image_base64:
                        # Это изображение - добавляем в список для отправки в vision модель
                        try:
                            file_name_lower = file_name.lower()
                            if not file_type or file_type == 'неизвестный тип':
                                if file_name_lower.endswith('.png'):
                                    file_type = 'image/png'
                                elif file_name_lower.endswith('.jpg') or file_name_lower.endswith('.jpeg'):
                                    file_type = 'image/jpeg'
                                elif file_name_lower.endswith('.gif'):
                                    file_type = 'image/gif'
                                elif file_name_lower.endswith('.webp'):
                                    file_type = 'image/webp'
                                else:
                                    file_type = 'image/jpeg'
                            
                            image_base64_clean = image_base64.strip().replace('\n', '').replace('\r', '').replace(' ', '')
                            
                            image_files.append({
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{file_type};base64,{image_base64_clean}"
                                }
                            })
                            file_contents.append(f"Файл {i+1} ({file_name}): {extracted_text}")
                        except Exception as e:
                            logger.error(f"Ошибка при подготовке изображения '{file_name}': {str(e)}", exc_info=True)
                            file_contents.append(f"Файл {i+1} ({file_name}): [Ошибка при подготовке изображения: {str(e)}]")
                    else:
                        # Это текстовый файл или документ
                        if extracted_text and not extracted_text.startswith('[') and not extracted_text.startswith('Изображение'):
                            text_preview = extracted_text[:1500]  # Ограничиваем до 1500 символов
                            if len(extracted_text) > 1500:
                                text_preview += "\n... [текст обрезан, показаны первые 1500 символов]"
                            file_contents.append(f"Файл {i+1} ({file_name}):\n{text_preview}")
                        else:
                            file_contents.append(f"Файл {i+1} ({file_name}): {extracted_text}")
                except Exception as e:
                    logger.error(f"Критическая ошибка при обработке файла: {str(e)}", exc_info=True)
                    file_name = file.get('name', f'Файл {i+1}') if isinstance(file, dict) else f'Файл {i+1}'
                    file_contents.append(f"Файл {i+1} ({file_name}): [Критическая ошибка при обработке: {str(e)}]")
        
        # Добавляем историю чата (последние 5 сообщений, ограничиваем размер)
        for msg in chat_history[-5:]:
            msg_text = msg.get('text', '')
            if len(msg_text) > 300:  # Обрезаем если больше 300 символов
                msg_text = msg_text[:300] + "... [обрезано]"
            role = "user" if msg.get('isUser') else "assistant"
            if role != "system":
                messages.append({
                    "role": role,
                    "content": msg_text
                })
        
        # Формируем финальное сообщение пользователя
        final_content_parts = []
        
        if file_contents:
            file_info = "Пользователь прикрепил файлы:\n\n" + "\n\n".join(file_contents)
            if len(file_info) > 2000:  # Обрезаем если слишком длинно
                file_info = file_info[:2000] + "\n... [информация о файлах обрезана]"
            final_content_parts.append({"type": "text", "text": file_info})
        
        if message:
            message_limited = message[:800] if len(message) > 800 else message  # Обрезаем до 800 символов
            if len(message) > 800:
                message_limited += "... [сообщение обрезано]"
            
            if final_content_parts:
                final_content_parts.append({"type": "text", "text": f"\n\nСообщение: {message_limited}"})
            else:
                final_content_parts.append({"type": "text", "text": message_limited})
        
        if not final_content_parts and image_files:
            final_content_parts.append({"type": "text", "text": "Проанализируй прикрепленные изображения"})
        
        final_content_parts.extend(image_files)
        
        # Создаем финальное сообщение
        if final_content_parts:
            if len(final_content_parts) == 1 and final_content_parts[0]["type"] == "text":
                messages.append({
                    "role": "user",
                    "content": final_content_parts[0]["text"]
                })
            else:
                messages.append({
                    "role": "user",
                    "content": final_content_parts
                })
        
        # Проверяем системный промпт
        if not messages or messages[0].get("role") != "system":
            system_msg = None
            other_messages = []
            for msg in messages:
                if msg.get("role") == "system":
                    system_msg = msg
                else:
                    other_messages.append(msg)
            
            if system_msg:
                messages = [system_msg] + other_messages
            else:
                messages.insert(0, {
                    "role": "system",
                    "content": system_prompt
                })
        
        # Подготавливаем payload для запроса
        payload = {
            "model": OPENROUTER_MODEL,
            "messages": messages,
            "temperature": 0.5,
            "max_tokens": 1200,
            "top_p": 0.9,
            "frequency_penalty": 0.1,
            "stream": False
        }
        
        # Заголовки для OpenRouter
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://localhost",
        }
        
        logger.debug(f"Отправка запроса в OpenRouter: model={payload['model']}, messages_count={len(messages)}")
        
        # Отслеживание времени начала LLM обработки
        llm_start_time = timezone.now()
        
        # Отправляем запрос в OpenRouter
        try:
            response = requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=90)
        except requests.exceptions.RequestException as e:
            logger.error(f"Ошибка при отправке запроса в OpenRouter: {str(e)}")
            chat_request.status = ChatRequest.STATUS_FAILED
            chat_request.error = f'Ошибка подключения к OpenRouter: {str(e)}'
            chat_request.save()
            return
        
        if response.status_code == 200:
            result = response.json()
            ai_response = result.get('choices', [{}])[0].get('message', {}).get('content', 'Извините, не удалось получить ответ.')
            
            # Модерация ответа AI
            moderation_result = ContentModerator.check_ai_response(ai_response)
            response_blocked = not moderation_result['allowed']
            if not moderation_result['allowed']:
                logger.warning(f"Ответ AI заблокирован модератором: {moderation_result['reason']}")
                ai_response = "Извините, я не могу предоставить ответ на этот запрос. Пожалуйста, переформулируйте вопрос в рамках делового общения."
            else:
                ai_response = moderation_result['filtered_response']
            
            # Обрабатываем действия (CREATE_EVENT, UPDATE_EVENT, DELETE_EVENT, DELETE_DOCUMENT, RENAME_DOCUMENT, SEND_SUPPORT_MESSAGE)
            action_result = None
            
            # СНАЧАЛА обрабатываем простые текстовые команды (они имеют приоритет)
            if 'DELETE_EVENT:' in ai_response:
                parts = ai_response.split('DELETE_EVENT:')
                if len(parts) > 1:
                    event_identifier = parts[1].strip().strip('"').strip("'").strip()
                    # Убираем возможные переносы строк и лишние символы
                    event_identifier = event_identifier.split('\n')[0].split('|')[0].strip()
                    calendar_events = user_data.get('calendarEvents', [])
                    
                    logger.info(f"Поиск события для удаления: '{event_identifier}'")
                    logger.info(f"Всего событий в календаре: {len(calendar_events)}")
                    
                    # Используем умный поиск событий
                    event_id = find_event_smart(calendar_events, event_identifier)
                    
                    if event_id:
                        action_result = {
                            'action': 'delete_event',
                            'id': event_id
                        }
                        logger.info(f"Команда удаления события создана: {action_result}")
                        ai_response = parts[0].strip()
                    else:
                        logger.warning(f"Событие не найдено для удаления: '{event_identifier}'")
                        logger.warning(f"Доступные события: {[{'id': e.get('id'), 'title': e.get('title')} for e in calendar_events]}")
            
            if not action_result and 'DELETE_DOCUMENT:' in ai_response:
                parts = ai_response.split('DELETE_DOCUMENT:')
                if len(parts) > 1:
                    doc_identifier = parts[1].strip().strip('"').strip("'").strip()
                    documents = user_data.get('documents', [])
                    doc_id = None
                    
                    # Ищем по ID
                    for doc in documents:
                        if str(doc.get('id', '')) == doc_identifier:
                            doc_id = str(doc.get('id', ''))
                            break
                    
                    # Ищем по названию
                    if not doc_id:
                        doc_identifier_lower = doc_identifier.lower().strip()
                        for doc in documents:
                            doc_name = doc.get('name', '').lower().strip()
                            if doc_name == doc_identifier_lower or doc_identifier_lower in doc_name:
                                doc_id = str(doc.get('id', ''))
                                break
                    
                    if doc_id:
                        action_result = {
                            'action': 'delete_document',
                            'id': doc_id
                        }
                        ai_response = parts[0].strip()
            
            if not action_result and 'RENAME_DOCUMENT:' in ai_response:
                parts = ai_response.split('RENAME_DOCUMENT:')
                if len(parts) > 1:
                    rename_data = parts[1].strip()
                    # Формат: старое_название|новое_название
                    if '|' in rename_data:
                        rename_parts = rename_data.split('|')
                        if len(rename_parts) >= 2:
                            doc_identifier = rename_parts[0].strip().strip('"').strip("'").strip()
                            new_name = rename_parts[1].strip().strip('"').strip("'").strip()
                            documents = user_data.get('documents', [])
                            doc_id = None
                            
                            # Ищем по ID или названию
                            for doc in documents:
                                if str(doc.get('id', '')) == doc_identifier:
                                    doc_id = str(doc.get('id', ''))
                                    break
                            
                            if not doc_id:
                                doc_identifier_lower = doc_identifier.lower().strip()
                                for doc in documents:
                                    doc_name = doc.get('name', '').lower().strip()
                                    if doc_name == doc_identifier_lower or doc_identifier_lower in doc_name:
                                        doc_id = str(doc.get('id', ''))
                                        break
                            
                            if doc_id and new_name:
                                action_result = {
                                    'action': 'rename_document',
                                    'id': doc_id,
                                    'name': new_name
                                }
                                ai_response = parts[0].strip()
            
            if not action_result and 'SEND_SUPPORT_MESSAGE:' in ai_response:
                parts = ai_response.split('SEND_SUPPORT_MESSAGE:')
                if len(parts) > 1:
                    message_data = parts[1].strip()
                    # Формат: тема|сообщение
                    if '|' in message_data:
                        message_parts = message_data.split('|')
                        if len(message_parts) >= 2:
                            subject = message_parts[0].strip().strip('"').strip("'").strip()
                            message = message_parts[1].strip().strip('"').strip("'").strip()
                            if subject and message:
                                action_result = {
                                    'action': 'send_support_message',
                                    'subject': subject,
                                    'message': message
                                }
                                ai_response = parts[0].strip()
            
            if not action_result and 'CREATE_EVENT:' in ai_response:
                parts = ai_response.split('CREATE_EVENT:')
                if len(parts) > 1:
                    event_data_raw = parts[1].strip()
                    # Убираем возможные переносы строк и лишние символы
                    event_data_raw = event_data_raw.split('\n')[0].strip()
                    # Разделяем по |
                    event_data = []
                    current_part = ''
                    for char in event_data_raw:
                        if char == '|':
                            event_data.append(current_part.strip())
                            current_part = ''
                        else:
                            current_part += char
                    if current_part:
                        event_data.append(current_part.strip())
                    
                    logger.info(f"Обработка CREATE_EVENT: найдено {len(event_data)} частей")
                    logger.info(f"Данные события: {event_data}")
                    
                    if len(event_data) >= 2:
                        title = event_data[0].strip().strip('"').strip("'").strip()
                        date_str = event_data[1].strip().strip('"').strip("'").strip()
                        description = event_data[2].strip().strip('"').strip("'").strip() if len(event_data) > 2 else ''
                        
                        # Нормализуем дату
                        # Если дата в формате YYYY-MM-DDTHH:mm, оставляем как есть
                        # Если только дата, добавляем время
                        if date_str and 'T' not in date_str:
                            if ' ' in date_str:
                                date_str = date_str.replace(' ', 'T')
                            else:
                                date_str = date_str + 'T12:00'
                        
                        action_result = {
                            'action': 'create_event',
                            'title': title,
                            'date': date_str,
                            'description': description
                        }
                        logger.info(f"Создано действие create_event: {action_result}")
                        ai_response = parts[0].strip()
                    else:
                        logger.warning(f"CREATE_EVENT: недостаточно данных. Найдено частей: {len(event_data)}")
            
            # ПОТОМ обрабатываем JSON команды (если простые команды не сработали)
            if not action_result:
                # Ищем все JSON команды в ответе
                json_commands = []
                json_patterns = [
                    r'\{\s*"action"\s*:\s*"[^"]+"[^}]*\}',
                    r'\{[^{}]*"action"\s*:\s*"[^"]+"[^{}]*\}',
                ]
                
                for pattern in json_patterns:
                    matches = re.finditer(pattern, ai_response, re.DOTALL | re.IGNORECASE)
                    for match in matches:
                        try:
                            cmd = json.loads(match.group(0))
                            if cmd.get('action') in ['UPDATE_EVENT', 'DELETE_EVENT', 'DELETE_DOCUMENT', 'RENAME_DOCUMENT', 'SEND_SUPPORT_MESSAGE']:
                                json_commands.append((match.group(0), cmd))
                        except:
                            continue
                
                # Обрабатываем команды по порядку
                for json_str, cmd_data in json_commands:
                    action_type = cmd_data.get('action')
                    
                    if action_type == 'DELETE_EVENT':
                        event_identifier = cmd_data.get('event', '').strip()
                        calendar_events = user_data.get('calendarEvents', [])
                        
                        # Используем умный поиск событий
                        event_id = find_event_smart(calendar_events, event_identifier)
                        
                        if event_id:
                            action_result = {
                                'action': 'delete_event',
                                'id': event_id
                            }
                            ai_response = ai_response.replace(json_str, '').strip()
                    
                    elif action_type == 'DELETE_DOCUMENT':
                        doc_identifier = cmd_data.get('document', '').strip()
                        documents = user_data.get('documents', [])
                        doc_id = None
                        
                        # Ищем по ID
                        for doc in documents:
                            if str(doc.get('id', '')) == doc_identifier:
                                doc_id = str(doc.get('id', ''))
                                break
                        
                        # Ищем по названию
                        if not doc_id:
                            doc_identifier_lower = doc_identifier.lower().strip()
                            # Сначала точное совпадение
                            for doc in documents:
                                doc_name = doc.get('name', '').lower().strip()
                                if doc_name == doc_identifier_lower:
                                    doc_id = str(doc.get('id', ''))
                                    break
                            
                            # Если не нашли, ищем частичное совпадение
                            if not doc_id:
                                for doc in documents:
                                    doc_name = doc.get('name', '').lower().strip()
                                    if doc_identifier_lower in doc_name or doc_name in doc_identifier_lower:
                                        doc_id = str(doc.get('id', ''))
                                        break
                        
                        if doc_id:
                            action_result = {
                                'action': 'delete_document',
                                'id': doc_id
                            }
                            ai_response = ai_response.replace(json_str, '').strip()
                    
                    elif action_type == 'RENAME_DOCUMENT':
                        doc_identifier = cmd_data.get('document', '').strip()
                        new_name = cmd_data.get('name', '').strip()
                        documents = user_data.get('documents', [])
                        doc_id = None
                        
                        for doc in documents:
                            if str(doc.get('id', '')) == doc_identifier:
                                doc_id = str(doc.get('id', ''))
                                break
                        
                        if not doc_id:
                            doc_identifier_lower = doc_identifier.lower().strip()
                            for doc in documents:
                                doc_name = doc.get('name', '').lower().strip()
                                if doc_name == doc_identifier_lower:
                                    doc_id = str(doc.get('id', ''))
                                    break
                        
                        if doc_id and new_name:
                            action_result = {
                                'action': 'rename_document',
                                'id': doc_id,
                                'name': new_name
                            }
                            ai_response = ai_response.replace(json_str, '').strip()
                    
                    elif action_type == 'SEND_SUPPORT_MESSAGE':
                        subject = cmd_data.get('subject', '').strip()
                        message = cmd_data.get('message', '').strip()
                        if subject and message:
                            action_result = {
                                'action': 'send_support_message',
                                'subject': subject,
                                'message': message
                            }
                            ai_response = ai_response.replace(json_str, '').strip()
            
            # СНАЧАЛА обрабатываем простой текстовый формат UPDATE_EVENT (более надежный)
            if not action_result and 'UPDATE_EVENT:' in ai_response:
                parts = ai_response.split('UPDATE_EVENT:')
                if len(parts) > 1:
                    event_data_raw = parts[1].strip()
                    # Убираем возможные переносы строк и лишние символы
                    event_data_raw = event_data_raw.split('\n')[0].strip()
                    # Разделяем по |
                    event_data = []
                    current_part = ''
                    for char in event_data_raw:
                        if char == '|':
                            event_data.append(current_part.strip())
                            current_part = ''
                        else:
                            current_part += char
                    if current_part:
                        event_data.append(current_part.strip())
                    
                    logger.info(f"Обработка UPDATE_EVENT (текстовый формат): найдено {len(event_data)} частей")
                    logger.info(f"Данные события: {event_data}")
                    
                    if len(event_data) >= 1:
                        event_identifier = event_data[0].strip().strip('"').strip("'").strip()
                        calendar_events = user_data.get('calendarEvents', [])
                        
                        # Используем умный поиск событий
                        event_id = find_event_smart(calendar_events, event_identifier)
                        
                        if event_id:
                            # Формат: UPDATE_EVENT: старое_название|новое_название|новая_дата|новое_описание
                            new_title = event_data[1].strip().strip('"').strip("'").strip() if len(event_data) > 1 and event_data[1].strip() else None
                            new_date = event_data[2].strip().strip('"').strip("'").strip() if len(event_data) > 2 and event_data[2].strip() else None
                            new_description = event_data[3].strip().strip('"').strip("'").strip() if len(event_data) > 3 and event_data[3].strip() else None
                            
                            # Нормализуем дату, если указана
                            if new_date and 'T' not in new_date:
                                if ' ' in new_date:
                                    new_date = new_date.replace(' ', 'T')
                                else:
                                    new_date = new_date + 'T12:00'
                            
                            action_result = {
                                'action': 'update_event',
                                'id': event_id,
                                'title': new_title,
                                'date': new_date,
                                'description': new_description
                            }
                            logger.info(f"Создано действие update_event (текстовый формат): {action_result}")
                            ai_response = parts[0].strip()
            
            # Затем обрабатываем JSON формат UPDATE_EVENT (для обратной совместимости)
            json_found = False
            json_str = None
            update_data = None
            
            if not action_result:
                json_patterns = [
                    r'\{\s*"action"\s*:\s*"UPDATE_EVENT"[^}]*\}',
                    r'\{[^{}]*"action"\s*:\s*"UPDATE_EVENT"[^{}]*\}',
                ]
                
                for pattern in json_patterns:
                    json_match = re.search(pattern, ai_response, re.DOTALL | re.IGNORECASE)
                    if json_match:
                        json_str = json_match.group(0)
                        try:
                            update_data = json.loads(json_str)
                            if update_data.get('action') == 'UPDATE_EVENT':
                                json_found = True
                                break
                        except:
                            continue
                
                if not json_found:
                    lines = ai_response.split('\n')
                    for line in lines:
                        line = line.strip()
                        if 'UPDATE_EVENT' in line and '{' in line and '}' in line:
                            try:
                                update_data = json.loads(line)
                                if update_data.get('action') == 'UPDATE_EVENT':
                                    json_str = line
                                    json_found = True
                                    break
                            except:
                                continue
            
            if json_found and update_data and not action_result:
                if update_data.get('action') == 'UPDATE_EVENT':
                    event_identifier = update_data.get('event', '').strip()
                    calendar_events = user_data.get('calendarEvents', [])
                    
                    logger.info(f"Поиск события для обновления: '{event_identifier}'")
                    logger.info(f"Всего событий в календаре: {len(calendar_events)}")
                    
                    # Используем умный поиск событий
                    event_id = find_event_smart(calendar_events, event_identifier)
                    
                    if not event_id and calendar_events:
                        event_id = str(calendar_events[-1].get('id', ''))
                        logger.warning(f"Событие не найдено, используем последнее: {event_id}")
                    
                    if event_id:
                        new_title = None
                        if 'title' in update_data:
                            title_value = update_data.get('title')
                            new_title = title_value.strip() if title_value else ''
                        
                        new_date = None
                        if 'date' in update_data:
                            new_date = update_data.get('date', '').strip() if update_data.get('date') else ''
                        
                        new_description = None
                        if 'description' in update_data:
                            desc_value = update_data.get('description', '')
                            if desc_value is not None:
                                new_description = desc_value.strip() if desc_value else ''
                            else:
                                new_description = ''
                        
                        action_result = {
                            'action': 'update_event',
                            'id': event_id,
                            'title': new_title,
                            'date': new_date,
                            'description': new_description
                        }
                        if json_str:
                            ai_response = ai_response.replace(json_str, '').strip()
            
            # Если не нашли JSON, пробуем старый формат UPDATE_EVENT:
            if not json_found and not action_result and 'UPDATE_EVENT:' in ai_response:
                parts = ai_response.split('UPDATE_EVENT:')
                if len(parts) > 1:
                    event_data_raw = parts[1].strip()
                    event_data = []
                    current_part = ''
                    for char in event_data_raw:
                        if char == '|':
                            event_data.append(current_part.strip())
                            current_part = ''
                        else:
                            current_part += char
                    if current_part:
                        event_data.append(current_part.strip())
                    
                    if len(event_data) >= 1:
                        event_identifier = event_data[0].strip()
                        calendar_events = user_data.get('calendarEvents', [])
                        
                        # Используем умный поиск событий
                        event_id = find_event_smart(calendar_events, event_identifier)
                        
                        if not event_id and calendar_events:
                            event_id = str(calendar_events[-1].get('id', ''))
                            logger.warning(f"Событие не найдено, используем последнее: {event_id}")
                        
                        if event_id:
                            new_title = event_data[1].strip() if len(event_data) > 1 and event_data[1].strip() else None
                            new_date = event_data[2].strip() if len(event_data) > 2 and event_data[2].strip() else None
                            new_description = event_data[3].strip() if len(event_data) > 3 and event_data[3].strip() else None
                            
                            action_result = {
                                'action': 'update_event',
                                'id': event_id,
                                'title': new_title,
                                'date': new_date,
                                'description': new_description
                            }
                        ai_response = parts[0].strip()
            
            # Дополнительная проверка: ищем команды в любом месте ответа (более гибкий поиск)
            if not action_result:
                # Ищем CREATE_EVENT в любом регистре и в любом месте
                create_patterns = [
                    r'CREATE_EVENT\s*:\s*([^\n]+)',
                    r'create_event\s*:\s*([^\n]+)',
                    r'Create_Event\s*:\s*([^\n]+)',
                ]
                for pattern in create_patterns:
                    match = re.search(pattern, ai_response, re.IGNORECASE)
                    if match:
                        event_data_raw = match.group(1).strip()
                        event_data = []
                        current_part = ''
                        for char in event_data_raw:
                            if char == '|':
                                event_data.append(current_part.strip())
                                current_part = ''
                            else:
                                current_part += char
                        if current_part:
                            event_data.append(current_part.strip())
                        
                        if len(event_data) >= 2:
                            title = event_data[0].strip().strip('"').strip("'").strip()
                            date_str = event_data[1].strip().strip('"').strip("'").strip()
                            description = event_data[2].strip().strip('"').strip("'").strip() if len(event_data) > 2 else ''
                            
                            if date_str and 'T' not in date_str:
                                if ' ' in date_str:
                                    date_str = date_str.replace(' ', 'T')
                                else:
                                    date_str = date_str + 'T12:00'
                            
                            action_result = {
                                'action': 'create_event',
                                'title': title,
                                'date': date_str,
                                'description': description
                            }
                            logger.info(f"Найдена команда CREATE_EVENT (гибкий поиск): {action_result}")
                            ai_response = ai_response.replace(match.group(0), '').strip()
                            break
                
                # Ищем DELETE_EVENT в любом месте
                delete_patterns = [
                    r'DELETE_EVENT\s*:\s*["\']?([^"\'\n]+)["\']?',
                    r'delete_event\s*:\s*["\']?([^"\'\n]+)["\']?',
                ]
                for pattern in delete_patterns:
                    match = re.search(pattern, ai_response, re.IGNORECASE)
                    if match:
                        event_identifier = match.group(1).strip().strip('"').strip("'").strip()
                        calendar_events = user_data.get('calendarEvents', [])
                        event_id = None
                        
                        for event in calendar_events:
                            if str(event.get('id', '')) == event_identifier:
                                event_id = str(event.get('id', ''))
                                break
                        
                        if not event_id:
                            event_identifier_lower = event_identifier.lower().strip()
                            for event in calendar_events:
                                event_title = event.get('title', '').lower().strip()
                                if event_title == event_identifier_lower or event_identifier_lower in event_title:
                                    event_id = str(event.get('id', ''))
                                    break
                        
                        if event_id:
                            action_result = {
                                'action': 'delete_event',
                                'id': event_id
                            }
                            logger.info(f"Найдена команда DELETE_EVENT (гибкий поиск): {action_result}")
                            ai_response = ai_response.replace(match.group(0), '').strip()
                            break
            
            # Сохраняем результат
            chat_request.status = ChatRequest.STATUS_COMPLETED
            chat_request.response = ai_response
            chat_request.action = action_result if action_result else {}
            chat_request.completed_at = timezone.now()
            chat_request.save()
            
            # Расчет метрик
            processing_end_time = timezone.now()
            processing_time = (processing_end_time - processing_start_time).total_seconds()
            llm_processing_time = None
            if llm_start_time:
                llm_processing_time = (processing_end_time - llm_start_time).total_seconds()
            
            # Определяем, использован ли контекст (если есть user_data)
            context_used = bool(user_data and len(user_data) > 0)
            
            # Определяем успешность действия
            action_success = None
            if action_result:
                # Действие считается успешным, если нет ошибки
                action_success = True  # Будет обновлено при фактическом выполнении действия
            
            # Создаем метрики для запроса
            try:
                MetricsCalculator.create_request_metrics(
                    chat_request=chat_request,
                    processing_time=processing_time,
                    llm_time=llm_processing_time,
                    has_action=bool(action_result),
                    action_success=action_success,
                    files_data=files,
                    message_blocked=message_blocked,
                    response_blocked=response_blocked,
                    context_used=context_used,
                    response_text=ai_response
                )
            except Exception as e:
                logger.error(f"Ошибка при создании метрик для запроса {request_id}: {str(e)}", exc_info=True)
            
            # Автоматический пересчет агрегированных метрик
            # Пересчитываем каждые 10 завершенных запросов или раз в час
            try:
                # Подсчитываем количество завершенных запросов за последний час
                one_hour_ago = timezone.now() - timedelta(hours=1)
                recent_completed = ChatRequest.objects.filter(
                    status=ChatRequest.STATUS_COMPLETED,
                    completed_at__gte=one_hour_ago
                ).count()
                
                # Пересчитываем метрики если:
                # 1. Прошло более 10 завершенных запросов с последнего расчета
                # 2. Или прошло более часа с последнего расчета
                last_metric = Metric.objects.order_by('-calculated_at').first()
                should_recalculate = False
                
                if last_metric:
                    time_since_last = timezone.now() - last_metric.calculated_at
                    # Если прошло более часа или более 10 запросов
                    if time_since_last.total_seconds() > 3600 or recent_completed % 10 == 0:
                        should_recalculate = True
                else:
                    # Если метрик еще нет, рассчитываем
                    should_recalculate = True
                
                if should_recalculate:
                    # Рассчитываем метрики за последние 7 дней в фоновом потоке
                    def calculate_metrics_async():
                        try:
                            period_end = timezone.now()
                            period_start = period_end - timedelta(days=7)
                            logger.info(f"Автоматический расчет метрик за период: {period_start} - {period_end}")
                            metrics = MetricsCalculator.calculate_all_metrics(period_start, period_end)
                            logger.info(f"Рассчитано {len(metrics)} метрик автоматически")
                        except Exception as e:
                            logger.error(f"Ошибка при автоматическом расчете метрик: {str(e)}", exc_info=True)
                    
                    # Запускаем в отдельном потоке, чтобы не блокировать обработку запроса
                    metric_thread = threading.Thread(target=calculate_metrics_async)
                    metric_thread.daemon = True
                    metric_thread.start()
                    
            except Exception as e:
                logger.error(f"Ошибка при автоматическом пересчете метрик: {str(e)}", exc_info=True)
            
            # Логируем действие для отладки
            if action_result:
                logger.info(f"✅ Действие сохранено в запрос: {action_result}")
            else:
                # Проверяем, содержит ли ответ команду CREATE_EVENT, которую не удалось распарсить
                if 'CREATE_EVENT' in ai_response:
                    logger.warning(f"⚠️ CREATE_EVENT найден в ответе, но не распознан. Ответ: {ai_response[:500]}")
                else:
                    logger.warning(f"⚠️ Действие не найдено в ответе AI. Ответ: {ai_response[:200]}")
            
            # Сохраняем в историю чатов (если есть email пользователя)
            try:
                user_email = user_data.get('email', '')
                if user_email:
                    # Получаем chat_id из истории сообщений или создаем новый
                    chat_id = None
                    if chat_history and len(chat_history) > 0:
                        # Пытаемся найти chat_id в метаданных
                        for msg in chat_history:
                            if isinstance(msg, dict) and 'chatId' in msg:
                                chat_id = msg.get('chatId')
                                break
                    
                    if not chat_id:
                        chat_id = f"chat_{int(time.time() * 1000)}"
                    
                    # Получаем или создаем запись истории
                    chat_history_obj, created = ChatHistory.objects.get_or_create(
                        user_email=user_email,
                        chat_id=chat_id,
                        defaults={
                            'title': 'Новый чат',
                            'messages': [],
                            'ai_actions': [],
                            'last_message_at': timezone.now()
                        }
                    )
                    
                    # Добавляем новое сообщение пользователя
                    user_message = {
                        'text': message,
                        'isUser': True,
                        'timestamp': timezone.now().isoformat(),
                        'files': [f.get('name', '') for f in files] if files else []
                    }
                    
                    # Добавляем ответ AI
                    ai_message = {
                        'text': ai_response,
                        'isUser': False,
                        'timestamp': timezone.now().isoformat(),
                        'action': action_result if action_result else None
                    }
                    
                    # Обновляем историю
                    if not chat_history_obj.messages:
                        chat_history_obj.messages = []
                    
                    chat_history_obj.messages.append(user_message)
                    chat_history_obj.messages.append(ai_message)
                    
                    # Добавляем действие AI в логи, если есть
                    if action_result:
                        action_log = {
                            'action': action_result.get('action', ''),
                            'data': action_result,
                            'timestamp': timezone.now().isoformat(),
                            'message': message,
                            'response': ai_response[:200]  # Первые 200 символов ответа
                        }
                        if not chat_history_obj.ai_actions:
                            chat_history_obj.ai_actions = []
                        chat_history_obj.ai_actions.append(action_log)
                        chat_history_obj.total_actions += 1
                    
                    # Обновляем статистику
                    chat_history_obj.total_messages = len(chat_history_obj.messages)
                    chat_history_obj.total_user_messages = sum(1 for m in chat_history_obj.messages if m.get('isUser', False))
                    chat_history_obj.total_ai_messages = sum(1 for m in chat_history_obj.messages if not m.get('isUser', False))
                    chat_history_obj.last_message_at = timezone.now()
                    chat_history_obj.updated_at = timezone.now()
                    
                    # Обновляем заголовок, если это первый чат
                    if created or not chat_history_obj.title or chat_history_obj.title == 'Новый чат':
                        if message:
                            title = message[:50] + ('...' if len(message) > 50 else '')
                            chat_history_obj.title = title
                    
                    chat_history_obj.save()
            except Exception as e:
                logger.error(f"Ошибка при сохранении истории чата: {str(e)}", exc_info=True)
            
        else:
            # Ошибка от OpenRouter
            error_details = f'Ошибка OpenRouter: {response.status_code}'
            try:
                error_data = response.json()
                if isinstance(error_data, dict):
                    error_message = error_data.get('error', {}).get('message', '') if isinstance(error_data.get('error'), dict) else str(error_data.get('error', ''))
                    if error_message:
                        error_details = f'Ошибка OpenRouter ({response.status_code}): {error_message}'
                    else:
                        error_details = f'Ошибка OpenRouter ({response.status_code}): {json.dumps(error_data, ensure_ascii=False)}'
                else:
                    error_details = f'Ошибка OpenRouter ({response.status_code}): {str(error_data)}'
            except:
                error_details = f'Ошибка OpenRouter ({response.status_code}): {response.text[:500]}'
            
            logger.error(f"Ошибка от OpenRouter: {error_details}")
            
            # Если ошибка 400, пробуем без изображений
            if response.status_code == 400 and image_files:
                try:
                    text_only_messages = []
                    for msg in messages:
                        if isinstance(msg.get('content'), list):
                            text_parts = [item for item in msg['content'] if item.get('type') == 'text']
                            if text_parts:
                                text_content = '\n'.join([item.get('text', '') for item in text_parts])
                                text_content += '\n\n[Примечание: изображения не были обработаны из-за ошибки формата]'
                                text_only_messages.append({
                                    "role": msg.get('role', 'user'),
                                    "content": text_content
                                })
                        else:
                            text_only_messages.append(msg)
                    
                    if not text_only_messages or text_only_messages[0].get("role") != "system":
                        text_only_messages.insert(0, {
                            "role": "system",
                            "content": system_prompt
                        })
                    
                    text_payload = {**payload, "messages": text_only_messages}
                    try:
                        text_response = requests.post(OPENROUTER_URL, headers=headers, json=text_payload, timeout=90)
                        if text_response.status_code == 200:
                            result = text_response.json()
                            ai_response = result.get('choices', [{}])[0].get('message', {}).get('content', '')
                            
                            # Модерация ответа AI
                            moderation_result = ContentModerator.check_ai_response(ai_response)
                            if not moderation_result['allowed']:
                                logger.warning(f"Ответ AI заблокирован модератором: {moderation_result['reason']}")
                                ai_response = "Извините, я не могу предоставить ответ на этот запрос. Пожалуйста, переформулируйте вопрос в рамках делового общения."
                            else:
                                ai_response = moderation_result['filtered_response']
                            
                            chat_request.status = ChatRequest.STATUS_COMPLETED
                            chat_request.response = ai_response + '\n\n[Примечание: изображения не были обработаны. Возможно, модель не поддерживает формат изображений или требуется другой формат.]'
                            chat_request.action = {}
                            chat_request.completed_at = timezone.now()
                            chat_request.save()
                            
                            # Расчет метрик для случая с изображениями
                            processing_end_time = timezone.now()
                            processing_time = (processing_end_time - processing_start_time).total_seconds()
                            llm_processing_time = None
                            if llm_start_time:
                                llm_processing_time = (processing_end_time - llm_start_time).total_seconds()
                            
                            try:
                                MetricsCalculator.create_request_metrics(
                                    chat_request=chat_request,
                                    processing_time=processing_time,
                                    llm_time=llm_processing_time,
                                    has_action=False,
                                    action_success=None,
                                    files_data=files,
                                    message_blocked=message_blocked,
                                    response_blocked=response_blocked,
                                    context_used=bool(user_data and len(user_data) > 0),
                                    response_text=ai_response
                                )
                            except Exception as e:
                                logger.error(f"Ошибка при создании метрик: {str(e)}")
                            
                            return
                    except Exception as e:
                        logger.error(f"Ошибка при обработке текстового запроса: {str(e)}")
                        pass
                except Exception as e:
                    logger.error(f"Ошибка при обработке изображений: {str(e)}")
                    pass
            
            chat_request.status = ChatRequest.STATUS_FAILED
            chat_request.error = error_details
            chat_request.save()
            
    except ChatRequest.DoesNotExist:
        logger.error(f"ChatRequest {request_id} не найден")
    except requests.exceptions.ConnectionError as e:
        error_msg = f'Не удалось подключиться к OpenRouter. Проверьте интернет-соединение и настройки API ключа.'
        logger.error(f"Ошибка подключения к OpenRouter: {str(e)}")
        try:
            chat_request = ChatRequest.objects.get(id=request_id)
            chat_request.status = ChatRequest.STATUS_FAILED
            chat_request.error = error_msg
            chat_request.save()
        except:
            pass
    except requests.exceptions.Timeout:
        error_msg = 'Превышено время ожидания ответа от OpenRouter (90 сек).'
        logger.error(error_msg)
        try:
            chat_request = ChatRequest.objects.get(id=request_id)
            chat_request.status = ChatRequest.STATUS_FAILED
            chat_request.error = error_msg
            chat_request.save()
        except:
            pass
    except Exception as e:
        logger.error(f"Ошибка при обработке запроса {request_id}: {str(e)}", exc_info=True)
        try:
            chat_request = ChatRequest.objects.get(id=request_id)
            chat_request.status = ChatRequest.STATUS_FAILED
            chat_request.error = f'Ошибка обработки: {str(e)}'
            chat_request.save()
        except:
            pass


@csrf_exempt
@require_http_methods(["POST"])
def chat_api(request):
    """API endpoint для обработки сообщений чата через OpenRouter (асинхронно)"""
    try:
        # Проверяем размер запроса
        content_length = request.META.get('CONTENT_LENGTH', 0)
        if content_length:
            content_length = int(content_length)
            max_size = getattr(settings, 'DATA_UPLOAD_MAX_MEMORY_SIZE', 2621440)  # По умолчанию 2.5MB
            if content_length > max_size:
                return JsonResponse({
                    'success': False,
                    'error': f'Размер данных ({content_length / 1024 / 1024:.2f} MB) превышает максимально допустимый ({max_size / 1024 / 1024:.2f} MB). Пожалуйста, уменьшите размер файлов.'
                }, status=413)
        
        data = json.loads(request.body)
        message = data.get('message', '')
        chat_history = data.get('history', [])
        user_data = data.get('userData', {})  # Данные пользователя из localStorage
        files = data.get('files', [])  # Прикрепленные файлы
        
        # Модерация входящего сообщения
        message = ContentModerator.sanitize_message(message)
        moderation_result = ContentModerator.check_message(message)
        
        if not moderation_result['allowed']:
            logger.warning(f"Сообщение заблокировано модератором: {moderation_result['reason']}")
            return JsonResponse({
                'success': False,
                'error': f"Сообщение не может быть обработано: {moderation_result['reason']}",
                'error_code': 'CONTENT_MODERATION_FAILED'
            }, status=400)
        
        # Используем отфильтрованное сообщение
        message = moderation_result['filtered_message']
        
        # Создаем запрос в базе данных
        chat_request = ChatRequest.objects.create(
            message=message,
            chat_history=chat_history,
            user_data=user_data,
            files_data=files,
            status=ChatRequest.STATUS_PENDING
        )
        
        # Создаем запись UserActivity для отслеживания запроса
        try:
            user_email = user_data.get('email', '') or user_data.get('userEmail', '')
            user = None
            if user_email:
                try:
                    user = User.objects.get(email=user_email)
                except User.DoesNotExist:
                    pass
            
            ip_address = request.META.get('REMOTE_ADDR', '')
            user_agent = request.META.get('HTTP_USER_AGENT', '')
            
            UserActivity.objects.create(
                user=user,
                user_email=user_email if user_email else None,
                activity_type='chat_request',
                activity_data={
                    'message_length': len(message),
                    'has_files': bool(files and len(files) > 0),
                    'files_count': len(files) if files else 0,
                    'chat_history_length': len(chat_history) if chat_history else 0
                },
                chat_request=chat_request,
                ip_address=ip_address if ip_address else None,
                user_agent=user_agent
            )
        except Exception as e:
            logger.error(f"Ошибка при создании UserActivity для запроса: {str(e)}", exc_info=True)
            # Не прерываем обработку запроса из-за ошибки метрик
        
        # Запускаем обработку в фоновом потоке
        thread = threading.Thread(target=process_chat_request_async, args=(chat_request.id,))
        thread.daemon = True
        thread.start()
        
        # Сразу возвращаем ID запроса
        return JsonResponse({
            'success': True,
            'request_id': str(chat_request.id),
            'status': 'processing',
            'message': 'Запрос принят в обработку'
        })
        
    except json.JSONDecodeError as e:
        logger.error(f"Ошибка парсинга JSON: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': 'Неверный формат данных. Проверьте корректность JSON.',
            'error_code': 'INVALID_JSON'
        }, status=400)
    except ValueError as e:
        logger.error(f"Ошибка валидации данных: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка валидации: {str(e)}',
            'error_code': 'VALIDATION_ERROR'
        }, status=400)
    except Exception as e:
        logger.error(f"Ошибка при создании запроса: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': 'Внутренняя ошибка сервера. Попробуйте позже.',
            'error_code': 'INTERNAL_ERROR',
            'error_id': str(time.time())  # ID для отслеживания ошибки
        }, status=500)


@csrf_exempt
@require_http_methods(["GET"])
def get_chat_history(request):
    """API endpoint для получения списка истории чатов пользователя"""
    try:
        user_email = request.GET.get('email', '')
        if not user_email:
            return JsonResponse({
                'success': False,
                'error': 'Email не указан'
            }, status=400)
        
        # Получаем историю чатов пользователя
        chat_histories = ChatHistory.objects.filter(user_email=user_email).order_by('-last_message_at', '-created_at')[:100]
        
        history_list = []
        for chat in chat_histories:
            history_list.append({
                'id': str(chat.id),
                'chat_id': chat.chat_id,
                'title': chat.title,
                'total_messages': chat.total_messages,
                'total_user_messages': chat.total_user_messages,
                'total_ai_messages': chat.total_ai_messages,
                'total_actions': chat.total_actions,
                'created_at': chat.created_at.isoformat() if chat.created_at else None,
                'updated_at': chat.updated_at.isoformat() if chat.updated_at else None,
                'last_message_at': chat.last_message_at.isoformat() if chat.last_message_at else None
            })
        
        return JsonResponse({
            'success': True,
            'history': history_list,
            'total': len(history_list)
        })
    except Exception as e:
        logger.error(f"Ошибка при получении истории чатов: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["GET"])
def get_chat_history_detail(request, chat_id):
    """API endpoint для получения детальной информации о чате"""
    try:
        user_email = request.GET.get('email', '')
        if not user_email:
            return JsonResponse({
                'success': False,
                'error': 'Email не указан'
            }, status=400)
        
        try:
            chat = ChatHistory.objects.get(chat_id=chat_id, user_email=user_email)
        except ChatHistory.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'Чат не найден'
            }, status=404)
        
        return JsonResponse({
            'success': True,
            'chat': {
                'id': str(chat.id),
                'chat_id': chat.chat_id,
                'title': chat.title,
                'messages': chat.messages,
                'ai_actions': chat.ai_actions,
                'total_messages': chat.total_messages,
                'total_user_messages': chat.total_user_messages,
                'total_ai_messages': chat.total_ai_messages,
                'total_actions': chat.total_actions,
                'created_at': chat.created_at.isoformat() if chat.created_at else None,
                'updated_at': chat.updated_at.isoformat() if chat.updated_at else None,
                'last_message_at': chat.last_message_at.isoformat() if chat.last_message_at else None
            }
        })
    except Exception as e:
        logger.error(f"Ошибка при получении деталей чата: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["GET"])
def export_chat_history(request, chat_id, format):
    """API endpoint для экспорта истории чата в различных форматах"""
    try:
        user_email = request.GET.get('email', '')
        
        # Ищем чат по chat_id (строка, например "chat_1764598137988"), email опционален
        try:
            if user_email:
                chat = ChatHistory.objects.get(chat_id=chat_id, user_email=user_email)
            else:
                chat = ChatHistory.objects.get(chat_id=chat_id)
        except ChatHistory.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'Чат не найден'
            }, status=404)
        
        format = format.lower()
        
        if format == 'json':
            response = HttpResponse(
                json.dumps({
                    'title': chat.title,
                    'chat_id': chat.chat_id,
                    'created_at': chat.created_at.isoformat() if chat.created_at else None,
                    'updated_at': chat.updated_at.isoformat() if chat.updated_at else None,
                    'last_message_at': chat.last_message_at.isoformat() if chat.last_message_at else None,
                    'statistics': {
                        'total_messages': chat.total_messages,
                        'total_user_messages': chat.total_user_messages,
                        'total_ai_messages': chat.total_ai_messages,
                        'total_actions': chat.total_actions
                    },
                    'messages': chat.messages,
                    'ai_actions': chat.ai_actions
                }, ensure_ascii=False, indent=2),
                content_type='application/json; charset=utf-8'
            )
            response['Content-Disposition'] = f'attachment; filename="chat_{chat.chat_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json"'
            return response
        
        elif format == 'docx':
            try:
                # Создаем документ Word
                doc = Document()
                
                # Настройка шрифта по умолчанию
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Calibri'
                font.size = Pt(11)
                
                # Заголовок документа
                title = doc.add_heading(chat.title or 'Новый чат', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title.runs[0]
                title_run.font.size = Pt(24)
                title_run.font.bold = True
                title_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Информация о чате
                doc.add_paragraph().add_run('Информация о чате').bold = True
                info_para = doc.add_paragraph()
                info_para.add_run(f'ID чата: ').bold = True
                info_para.add_run(str(chat.chat_id))
                info_para = doc.add_paragraph()
                info_para.add_run(f'Создан: ').bold = True
                info_para.add_run(chat.created_at.strftime('%Y-%m-%d %H:%M:%S') if chat.created_at else 'Неизвестно')
                info_para = doc.add_paragraph()
                info_para.add_run(f'Обновлен: ').bold = True
                info_para.add_run(chat.updated_at.strftime('%Y-%m-%d %H:%M:%S') if chat.updated_at else 'Неизвестно')
                info_para = doc.add_paragraph()
                info_para.add_run(f'Последнее сообщение: ').bold = True
                info_para.add_run(chat.last_message_at.strftime('%Y-%m-%d %H:%M:%S') if chat.last_message_at else 'Неизвестно')
                
                # Статистика
                doc.add_paragraph()  # Пустая строка
                stats_heading = doc.add_paragraph()
                stats_heading.add_run('Статистика').bold = True
                stats_heading.runs[0].font.size = Pt(14)
                
                stats_para = doc.add_paragraph()
                stats_para.add_run(f'Всего сообщений: ').bold = True
                stats_para.add_run(str(chat.total_messages))
                stats_para = doc.add_paragraph()
                stats_para.add_run(f'Сообщений пользователя: ').bold = True
                stats_para.add_run(str(chat.total_user_messages))
                stats_para = doc.add_paragraph()
                stats_para.add_run(f'Сообщений AI: ').bold = True
                stats_para.add_run(str(chat.total_ai_messages))
                stats_para = doc.add_paragraph()
                stats_para.add_run(f'Действий AI: ').bold = True
                stats_para.add_run(str(chat.total_actions))
                
                # Разделитель
                doc.add_paragraph()  # Пустая строка
                separator = doc.add_paragraph('─' * 80)
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                separator.runs[0].font.color.rgb = RGBColor(200, 200, 200)
                doc.add_paragraph()  # Пустая строка
                
                # Сообщения
                messages_heading = doc.add_heading('История сообщений', level=1)
                messages_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                
                # Проверяем, что messages существует и является списком
                messages_list = chat.messages if chat.messages and isinstance(chat.messages, list) else []
                
                for msg in messages_list:
                    timestamp = msg.get('timestamp', '')
                    if timestamp:
                        try:
                            dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                            timestamp_str = dt.strftime('%Y-%m-%d %H:%M:%S')
                        except:
                            timestamp_str = timestamp
                    else:
                        timestamp_str = 'Неизвестно'
                    
                    is_user = msg.get('isUser', False)
                    role = "Пользователь" if is_user else "AI-Помощник"
                    text = msg.get('text', '')
                    files = msg.get('files', [])
                    
                    # Заголовок сообщения
                    msg_heading = doc.add_paragraph()
                    msg_heading.add_run(f'[{timestamp_str}] ').font.color.rgb = RGBColor(128, 128, 128)
                    role_run = msg_heading.add_run(role)
                    role_run.bold = True
                    if is_user:
                        role_run.font.color.rgb = RGBColor(0, 102, 204)  # Синий для пользователя
                    else:
                        role_run.font.color.rgb = RGBColor(220, 20, 60)  # Красный для AI
                    
                    # Текст сообщения
                    text_content = str(text) if text else '(пустое сообщение)'
                    text_para = doc.add_paragraph(text_content)
                    text_para.style = 'List Paragraph'
                    text_para.paragraph_format.left_indent = Inches(0.5)
                    text_para.paragraph_format.space_after = Pt(6)
                    
                    # Прикрепленные файлы
                    if files:
                        files_para = doc.add_paragraph()
                        files_para.add_run('Прикрепленные файлы: ').font.color.rgb = RGBColor(100, 100, 100)
                        files_list = ', '.join([str(f) if isinstance(f, (str, int)) else str(f.get('name', f)) if isinstance(f, dict) else str(f) for f in files])
                        files_para.add_run(files_list).italic = True
                        files_para.paragraph_format.left_indent = Inches(0.5)
                    
                    # Действие AI
                    action = msg.get('action')
                    if action:
                        action_para = doc.add_paragraph()
                        action_para.add_run('Действие AI: ').font.color.rgb = RGBColor(100, 100, 100)
                        action_text = action.get('action', '') if isinstance(action, dict) else str(action)
                        action_para.add_run(action_text).italic = True
                        action_para.paragraph_format.left_indent = Inches(0.5)
                    
                    # Разделитель между сообщениями
                    doc.add_paragraph()  # Пустая строка
                    msg_separator = doc.add_paragraph('·' * 80)
                    msg_separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    msg_separator.runs[0].font.color.rgb = RGBColor(230, 230, 230)
                    doc.add_paragraph()  # Пустая строка
                
                # Действия AI
                ai_actions_list = chat.ai_actions if chat.ai_actions and isinstance(chat.ai_actions, list) else []
                if ai_actions_list:
                    doc.add_paragraph()  # Пустая строка
                    separator = doc.add_paragraph('═' * 80)
                    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    separator.runs[0].font.color.rgb = RGBColor(200, 200, 200)
                    doc.add_paragraph()  # Пустая строка
                    
                    actions_heading = doc.add_heading('Действия AI', level=1)
                    actions_heading.runs[0].font.color.rgb = RGBColor(220, 20, 60)
                    
                    for action in ai_actions_list:
                        timestamp = action.get('timestamp', '')
                        if timestamp:
                            try:
                                dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                                timestamp_str = dt.strftime('%Y-%m-%d %H:%M:%S')
                            except:
                                timestamp_str = timestamp
                        else:
                            timestamp_str = 'Неизвестно'
                        
                        action_heading = doc.add_paragraph()
                        action_heading.add_run(f'[{timestamp_str}] ').font.color.rgb = RGBColor(128, 128, 128)
                        action_heading.add_run(action.get('action', '')).bold = True
                        action_heading.runs[-1].font.color.rgb = RGBColor(220, 20, 60)
                        
                        if action.get('data'):
                            data_para = doc.add_paragraph()
                            data_para.add_run('Данные: ').bold = True
                            data_para.add_run(json.dumps(action.get('data', {}), ensure_ascii=False, indent=2))
                            data_para.paragraph_format.left_indent = Inches(0.5)
                            data_para.style = 'List Paragraph'
                        
                        if action.get('message'):
                            msg_para = doc.add_paragraph()
                            msg_para.add_run('Запрос: ').bold = True
                            msg_para.add_run(action.get('message', ''))
                            msg_para.paragraph_format.left_indent = Inches(0.5)
                            msg_para.style = 'List Paragraph'
                        
                        if action.get('response'):
                            resp_para = doc.add_paragraph()
                            resp_para.add_run('Ответ: ').bold = True
                            resp_para.add_run(action.get('response', ''))
                            resp_para.paragraph_format.left_indent = Inches(0.5)
                            resp_para.style = 'List Paragraph'
                        
                        doc.add_paragraph()  # Пустая строка
            
                # Сохраняем документ в BytesIO
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                response = HttpResponse(
                    doc_io.read(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="chat_{chat.chat_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
                return response
            except Exception as e:
                logger.error(f"Ошибка при создании DOCX файла: {str(e)}", exc_info=True)
                return JsonResponse({
                    'success': False,
                    'error': f'Ошибка при создании документа: {str(e)}'
                }, status=500)
        
        elif format == 'csv':
            response = HttpResponse(content_type='text/csv; charset=utf-8')
            response['Content-Disposition'] = f'attachment; filename="chat_{chat.chat_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv"'
            
            writer = csv.writer(response)
            writer.writerow(['Время', 'Роль', 'Сообщение', 'Файлы', 'Действие AI'])
            
            for msg in chat.messages:
                timestamp = msg.get('timestamp', '')
                if timestamp:
                    try:
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        timestamp_str = dt.strftime('%Y-%m-%d %H:%M:%S')
                    except:
                        timestamp_str = timestamp
                else:
                    timestamp_str = 'Неизвестно'
                
                role = "Пользователь" if msg.get('isUser', False) else "AI"
                text = msg.get('text', '').replace('\n', ' ').replace('\r', ' ')
                files = ', '.join(msg.get('files', []))
                action = msg.get('action', {}).get('action', '') if msg.get('action') else ''
                
                writer.writerow([timestamp_str, role, text, files, action])
            
            return response
        
        else:
            return JsonResponse({
                'success': False,
                'error': f'Неподдерживаемый формат: {format}'
            }, status=400)
    
    except Exception as e:
        logger.error(f"Ошибка при экспорте истории чата: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def export_chat_docx_direct(request):
    """API endpoint для экспорта чата в DOCX напрямую из данных, без обращения к БД"""
    try:
        data = json.loads(request.body)
        chat_title = data.get('title', 'Новый чат')
        messages = data.get('messages', [])
        chat_id = data.get('chat_id', 'unknown')
        
        if not messages:
            return JsonResponse({
                'success': False,
                'error': 'Нет сообщений для экспорта'
            }, status=400)
        
        try:
            # Создаем документ Word
            doc = Document()
            
            # Настройка шрифта по умолчанию
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Заголовок документа
            title = doc.add_heading(chat_title or 'Новый чат', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Пустая строка после заголовка
            doc.add_paragraph()
            
            # Сообщения (без заголовка "История сообщений")
            
            for msg in messages:
                timestamp = msg.get('timestamp', '')
                if timestamp:
                    try:
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        timestamp_str = dt.strftime('%Y-%m-%d %H:%M:%S')
                    except:
                        timestamp_str = timestamp
                else:
                    timestamp_str = 'Неизвестно'
                
                is_user = msg.get('isUser', False)
                role = "Пользователь" if is_user else "AI-Помощник"
                text = msg.get('text', '')
                files = msg.get('files', [])
                
                # Заголовок сообщения
                msg_heading = doc.add_paragraph()
                msg_heading.add_run(f'[{timestamp_str}] ').font.color.rgb = RGBColor(128, 128, 128)
                role_run = msg_heading.add_run(role)
                role_run.bold = True
                if is_user:
                    role_run.font.color.rgb = RGBColor(0, 102, 204)  # Синий для пользователя
                else:
                    role_run.font.color.rgb = RGBColor(220, 20, 60)  # Красный для AI
                
                # Текст сообщения - проверяем на наличие таблиц
                text_content = str(text) if text else '(пустое сообщение)'
                
                # Парсим markdown таблицы
                table_pattern = r'(\|[^\n]+\|\n\|[-\s|:]+\|\n(?:\|[^\n]+\|\n?)+)'
                tables = re.finditer(table_pattern, text_content)
                
                # Разбиваем текст на части: до таблицы, таблицы, после таблицы
                last_pos = 0
                parts = []
                
                for match in tables:
                    # Текст до таблицы
                    if match.start() > last_pos:
                        text_before = text_content[last_pos:match.start()].strip()
                        if text_before:
                            parts.append(('text', text_before))
                    
                    # Таблица
                    table_text = match.group(0)
                    parts.append(('table', table_text))
                    last_pos = match.end()
                
                # Текст после последней таблицы
                if last_pos < len(text_content):
                    text_after = text_content[last_pos:].strip()
                    if text_after:
                        parts.append(('text', text_after))
                
                # Если таблиц не найдено, добавляем весь текст
                if not parts:
                    parts.append(('text', text_content))
                
                # Обрабатываем части
                for part_type, part_content in parts:
                    if part_type == 'table':
                        # Парсим markdown таблицу и создаем Word таблицу
                        lines = part_content.strip().split('\n')
                        if len(lines) >= 2:
                            # Первая строка - заголовок
                            header_cells = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
                            
                            # Вторая строка - разделитель с выравниванием
                            separator_line = lines[1]
                            separator_cells = [cell.strip() for cell in separator_line.split('|') if cell.strip()]
                            
                            # Определяем выравнивание
                            def get_alignment(separator):
                                if not separator:
                                    return WD_ALIGN_PARAGRAPH.LEFT
                                separator = separator.strip()
                                if separator.startswith(':') and separator.endswith(':'):
                                    return WD_ALIGN_PARAGRAPH.CENTER
                                elif separator.endswith(':'):
                                    return WD_ALIGN_PARAGRAPH.RIGHT
                                else:
                                    return WD_ALIGN_PARAGRAPH.LEFT
                            
                            alignments = [get_alignment(sep) for sep in separator_cells]
                            
                            # Создаем таблицу Word
                            num_cols = len(header_cells)
                            if num_cols > 0:
                                table = doc.add_table(rows=1, cols=num_cols)
                                table.style = 'Light Grid Accent 1'  # Красивый стиль таблицы
                                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                
                                # Заголовок
                                header_row = table.rows[0]
                                for i, cell_text in enumerate(header_cells):
                                    cell = header_row.cells[i]
                                    # Очищаем ячейку и добавляем текст с форматированием
                                    cell.text = ''
                                    paragraph = cell.paragraphs[0]
                                    paragraph.alignment = alignments[i] if i < len(alignments) else WD_ALIGN_PARAGRAPH.LEFT
                                    run = paragraph.add_run(cell_text)
                                    run.font.bold = True
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                    
                                    # Устанавливаем фон ячейки через XML (правильный способ)
                                    try:
                                        from docx.oxml import OxmlElement
                                        from docx.oxml.ns import qn
                                        
                                        tc_pr = cell._element.get_or_add_tcPr()
                                        shading = OxmlElement('w:shd')
                                        shading.set(qn('w:fill'), '4472C4')  # Синий цвет
                                        shading.set(qn('w:val'), 'clear')
                                        tc_pr.append(shading)
                                    except Exception:
                                        # Если не получилось установить цвет, просто оставляем без фона
                                        pass
                                
                                # Данные
                                for line_idx in range(2, len(lines)):
                                    line = lines[line_idx].strip()
                                    if not line or not line.startswith('|'):
                                        continue
                                    data_cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                                    if len(data_cells) > 0:
                                        row = table.add_row()
                                        for i, cell_text in enumerate(data_cells):
                                            if i < num_cols:
                                                cell = row.cells[i]
                                                cell.text = ''
                                                paragraph = cell.paragraphs[0]
                                                paragraph.alignment = alignments[i] if i < len(alignments) else WD_ALIGN_PARAGRAPH.LEFT
                                                run = paragraph.add_run(cell_text)
                                                run.font.size = Pt(10)
                                
                                # Добавляем отступ после таблицы
                                doc.add_paragraph()
                    else:
                        # Обычный текст
                        if part_content:
                            text_para = doc.add_paragraph(part_content)
                            text_para.style = 'List Paragraph'
                            text_para.paragraph_format.left_indent = Inches(0.5)
                            text_para.paragraph_format.space_after = Pt(6)
                
                # Прикрепленные файлы
                if files:
                    files_para = doc.add_paragraph()
                    files_para.add_run('Прикрепленные файлы: ').font.color.rgb = RGBColor(100, 100, 100)
                    files_list = ', '.join([
                        str(f.get('name', f)) if isinstance(f, dict) else str(f) 
                        for f in files
                    ])
                    files_para.add_run(files_list).italic = True
                    files_para.paragraph_format.left_indent = Inches(0.5)
                
                # Действие AI
                action = msg.get('action')
                if action:
                    action_para = doc.add_paragraph()
                    action_para.add_run('Действие AI: ').font.color.rgb = RGBColor(100, 100, 100)
                    action_text = action.get('action', '') if isinstance(action, dict) else str(action)
                    action_para.add_run(action_text).italic = True
                    action_para.paragraph_format.left_indent = Inches(0.5)
                
                # Разделитель между сообщениями
                doc.add_paragraph()  # Пустая строка
                msg_separator = doc.add_paragraph('·' * 80)
                msg_separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                msg_separator.runs[0].font.color.rgb = RGBColor(230, 230, 230)
                doc.add_paragraph()  # Пустая строка
            
            # Сохраняем документ в BytesIO
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            response = HttpResponse(
                doc_io.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="chat_{chat_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
            return response
            
        except Exception as e:
            logger.error(f"Ошибка при создании DOCX файла: {str(e)}", exc_info=True)
            return JsonResponse({
                'success': False,
                'error': f'Ошибка при создании документа: {str(e)}'
            }, status=500)
            
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Неверный формат JSON'
        }, status=400)
    except Exception as e:
        logger.error(f"Ошибка при экспорте чата в DOCX: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def edit_chat_message(request, chat_id):
    """API endpoint для редактирования сообщения в истории чата"""
    try:
        data = json.loads(request.body)
        user_email = data.get('email', '')
        message_index = data.get('message_index', -1)
        new_text = data.get('new_text', '')
        
        if not user_email:
            return JsonResponse({
                'success': False,
                'error': 'Email не указан'
            }, status=400)
        
        if message_index < 0:
            return JsonResponse({
                'success': False,
                'error': 'Индекс сообщения не указан'
            }, status=400)
        
        try:
            chat = ChatHistory.objects.get(chat_id=chat_id, user_email=user_email)
        except ChatHistory.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'Чат не найден'
            }, status=404)
        
        if message_index >= len(chat.messages):
            return JsonResponse({
                'success': False,
                'error': 'Индекс сообщения выходит за границы'
            }, status=400)
        
        # Редактируем сообщение
        message = chat.messages[message_index]
        message['text'] = new_text
        message['edited'] = True
        message['edited_at'] = timezone.now().isoformat()
        
        # Сохраняем изменения
        chat.messages[message_index] = message
        chat.updated_at = timezone.now()
        chat.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Сообщение успешно отредактировано',
            'updated_message': message
        })
    
    except Exception as e:
        logger.error(f"Ошибка при редактировании сообщения: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["GET"])
def chat_status(request, request_id):
    """API endpoint для проверки статуса запроса к AI"""
    try:
        chat_request = ChatRequest.objects.get(id=request_id)
        
        response_data = {
            'success': True,
            'request_id': str(chat_request.id),
            'status': chat_request.status,
            'created_at': chat_request.created_at.isoformat(),
            'updated_at': chat_request.updated_at.isoformat(),
        }
        
        if chat_request.status == ChatRequest.STATUS_COMPLETED:
            response_data['response'] = chat_request.response
            response_data['action'] = chat_request.action
            logger.info(f"✅ Отправка ответа клиенту: action={chat_request.action}")
            if chat_request.completed_at:
                response_data['completed_at'] = chat_request.completed_at.isoformat()
        elif chat_request.status == ChatRequest.STATUS_FAILED:
            response_data['error'] = chat_request.error
        
        return JsonResponse(response_data)
    except ChatRequest.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Запрос не найден'
        }, status=404)
    except Exception as e:
        logger.error(f"Ошибка при получении статуса запроса {request_id}: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка: {str(e)}'
        }, status=500)
        
        # Добавляем системный промпт с описанием возможностей (строгий деловой стиль)
        system_prompt = """Ты - профессиональный бизнес-ассистент. Работай СТРОГО в деловом стиле. 

КРИТИЧЕСКИ ВАЖНО - СОБЛЮДЕНИЕ ДЕЛОВОГО СТИЛЯ:
- ВСЕГДА отвечай только в деловом, формальном стиле
- НИКОГДА не переключайся на неформальный, дружелюбный или развлекательный тон
- НИКОГДА не рассказывай анекдоты, не шути, не используй развлекательный контент
- Используй только профессиональную деловую лексику
- Будь кратким, точным и по делу
- Избегай эмоциональных выражений, смайликов, восклицательных знаков (кроме важных уведомлений)
- Всегда сохраняй деловой тон, даже если пользователь задает неформальные вопросы

ТВОИ ВОЗМОЖНОСТИ:

1. АНАЛИЗИРОВАТЬ ДАННЫЕ:
   - ЧЕКИ: просматривать все операции, суммы, анализировать расходы и доходы, находить самые крупные операции
   - ИНВЕНТАРИЗАЦИЯ: просматривать товары, категории, количество, стоимость каждого товара, общую стоимость инвентаризации, анализировать складские остатки и их стоимость
     * КРИТИЧЕСКИ ВАЖНО: 
       - Используй НАЗВАНИЯ товаров (поле "name"), а НЕ ID товаров
       - Используй НАЗВАНИЯ папок (поле "name" папки), а НЕ ID папок
       - В контексте есть структура "СТРУКТУРА ИНВЕНТАРИЗАЦИИ" - там каждая ПАПКА с товарами внутри
       - Когда пользователь просит показать инвентаризацию, перечисляй каждую ПАПКУ с товарами внутри
     * Пример правильного ответа: 
       "ПАПКА 'Электроника':
        - ТОВАР 'Ноутбук': количество 10, цена 50,000 ₽, стоимость 500,000 ₽
        - ТОВАР 'Монитор': количество 5, цена 15,000 ₽, стоимость 75,000 ₽"
     * Пример НЕПРАВИЛЬНОГО ответа: "Товар 1764515968195.84" - НЕ используй ID!
     * Всегда называй товары по их названиям из поля "ТОВАР", папки по названиям из поля "ПАПКА"
   - СОТРУДНИКИ: просматривать список сотрудников, их должности, зарплаты, рассчитывать общий фонд оплаты труда
   - КАЛЕНДАРЬ: просматривать события, планировать встречи, напоминать о предстоящих событиях
   - НАЛОГИ: проверять задолженность по каждому налогу отдельно и общую сумму
   - КОММУНАЛЬНЫЕ УСЛУГИ: проверять задолженность по каждой услуге и общую сумму
   - ДОКУМЕНТЫ: работать с загруженными файлами, анализировать их содержимое (PDF, DOCX, XLSX, TXT, изображения и другие форматы)
   - БАЛАНСЫ: знать балансы обоих счетов пользователя

2. ВЫПОЛНЯТЬ ДЕЙСТВИЯ:
   - Создавать события в календаре: КРИТИЧЕСКИ ВАЖНО - когда пользователь просит запланировать встречу или событие, ОБЯЗАТЕЛЬНО создавай событие. Используй формат: CREATE_EVENT: название|дата в формате ISO (YYYY-MM-DDTHH:mm)|описание
     * ОБЯЗАТЕЛЬНО вычисляй точную дату из относительных формулировок:
       - "через неделю" → текущая дата + 7 дней
       - "ровно через неделю" → текущая дата + 7 дней
       - "через месяц" → текущая дата + 30 дней
       - "завтра" → текущая дата + 1 день
       - "послезавтра" → текущая дата + 2 дня
       - "через 3 дня" → текущая дата + 3 дня
       - "в следующую пятницу" → вычисли дату следующей пятницы
     * ОБЯЗАТЕЛЬНО извлекай время из запроса:
       - "в 15:00" → 15:00
       - "в 14:30" → 14:30
       - "в полдень" → 12:00
       - "в полночь" → 00:00
     * ПРИМЕРЫ ЕСТЕСТВЕННЫХ ЗАПРОСОВ:
       - "ровно через неделю у меня будет день рождения, буду отмечать в 15:00" → CREATE_EVENT: День рождения|2024-12-09T15:00|Отмечаю день рождения
       - "запланируй встречу с клиентом завтра в 14:00" → CREATE_EVENT: Встреча с клиентом|2024-12-03T14:00|Встреча с клиентом
       - "через 3 дня у меня совещание в 10:00 про проект" → CREATE_EVENT: Совещание про проект|2024-12-05T10:00|Совещание про проект
     * ВАЖНО: ВСЕГДА вычисляй точную дату и время, даже если пользователь использует относительные формулировки
     * ВАЖНО: После создания события ОБЯЗАТЕЛЬНО сообщи пользователю, что событие добавлено в календарь с указанием точной даты и времени
   - Редактировать события в календаре: КРИТИЧЕСКИ ВАЖНО - когда пользователь просит изменить или перенести встречу, ОБЯЗАТЕЛЬНО используй команду UPDATE_EVENT в формате JSON на отдельной строке:
     {"action": "UPDATE_EVENT", "event": "название_или_описание_события", "title": "новое_название", "date": "2024-12-25T15:00", "description": "новое_описание"}
     * ПОИСК СОБЫТИЙ: Можешь использовать не только точное название, но и:
       - Часть названия: "встреча с клиентом", "совещание"
       - Описание события: "то событие про проект", "встреча завтра"
       - Время: "встреча в 15:00", "совещание на завтра"
       - Система сама найдет подходящее событие по частичному совпадению
     * Если нужно изменить только название: {"action": "UPDATE_EVENT", "event": "часть_названия_или_описание", "title": "новое_название"}
     * Если нужно перенести встречу на другой день/время: {"action": "UPDATE_EVENT", "event": "часть_названия", "date": "2024-12-25T15:00"}
       - При переносе можно использовать относительные даты: "завтра", "послезавтра", "через неделю"
       - Система автоматически вычислит точную дату на основе текущей даты
     * Если нужно изменить описание/заметки/комментарий: {"action": "UPDATE_EVENT", "event": "часть_названия", "description": "новое_описание"}
     * ВАЖНО: При изменении описания ВСЕГДА указывай поле "description" с новым текстом описания
     * ВАЖНО: Описание может быть любым текстом, который пользователь просит добавить или изменить
     * Формат даты ОБЯЗАТЕЛЬНО: YYYY-MM-DDTHH:mm (например: 2024-12-25T15:00 для 25 декабря 2024 в 15:00)
     * ВАЖНО: При переносе встречи ВСЕГДА указывай поле "date" с полной датой и временем в формате ISO
     * ПРИМЕРЫ ЕСТЕСТВЕННЫХ ЗАПРОСОВ:
       - "Перенеси встречу с клиентом на завтра в 15:00" → {"action": "UPDATE_EVENT", "event": "встреча с клиентом", "date": "2024-12-26T15:00"}
       - "Измени название совещания на 'Планирование проекта'" → {"action": "UPDATE_EVENT", "event": "совещание", "title": "Планирование проекта"}
       - "Добавь описание к завтрашней встрече: обсудить бюджет" → {"action": "UPDATE_EVENT", "event": "завтра", "description": "обсудить бюджет"}
     * Команда должна быть на отдельной строке в начале или конце ответа
   - Удалять события календаря: когда пользователь просит удалить встречу или событие, используй простую команду:
     DELETE_EVENT: название_или_ID_события
     Пример: DELETE_EVENT: покатушки
   - Управлять документами:
     * Удалять документы: DELETE_DOCUMENT: название_или_ID_документа
       Пример: DELETE_DOCUMENT: отчет.pdf
     * Переименовывать документы: RENAME_DOCUMENT: старое_название|новое_название
       Пример: RENAME_DOCUMENT: старый_файл.docx|новый_файл.docx
   - Отправлять сообщения в поддержку: когда пользователь просит отправить сообщение в поддержку, используй команду:
     SEND_SUPPORT_MESSAGE: тема|текст_сообщения
     Пример: SEND_SUPPORT_MESSAGE: Проблема с входом|Не могу войти в систему
   - Анализировать файлы: если пользователь прикрепил файл, анализируй его содержимое:
     * Текстовые файлы (TXT, CSV, JSON, XML, HTML, MD и др.) - читай и анализируй текст
     * PDF документы - извлекай и анализируй текст со всех страниц
     * DOCX документы - извлекай и анализируй текст и таблицы
     * XLSX таблицы - извлекай и анализируй данные из всех листов
     * Изображения - описывай содержимое, анализируй объекты, текст на изображениях, диаграммы и графики

3. ОТВЕЧАТЬ НА ВОПРОСЫ (СТРОГО ДЕЛОВОЙ СТИЛЬ):
   - Отвечай кратко, точно, профессионально
   - Используй конкретные данные из контекста пользователя
   - Если данных нет, сообщи об этом формально
   - При анализе чеков указывай конкретные суммы и даты
   - При работе с инвентаризацией указывай конкретные позиции и количество
     * ВСЕГДА используй НАЗВАНИЯ товаров (поле "name"), НИКОГДА не используй ID товаров
     * ВСЕГДА используй НАЗВАНИЯ папок, НИКОГДА не используй ID папок
     * Когда пользователь просит показать инвентаризацию, структурируй ответ по ПАПКАМ:
       - Перечисляй каждую ПАПКУ отдельно
       - Внутри каждой ПАПКИ перечисляй все ТОВАРЫ с их данными
     * Пример правильного ответа: 
       "ПАПКА 'Электроника':
        - ТОВАР 'Ноутбук': количество 10, цена 50,000 ₽, стоимость 500,000 ₽
        - ТОВАР 'Монитор': количество 5, цена 15,000 ₽, стоимость 75,000 ₽
        ПАПКА 'Канцелярия':
        - ТОВАР 'Карандаш': количество 100, цена 50 ₽, стоимость 5,000 ₽"
     * Пример НЕПРАВИЛЬНОГО ответа: "Товар 1764515968195.84" - это ID, а не название, НЕ используй его!
     * Всегда используй структуру из "СТРУКТУРА ИНВЕНТАРИЗАЦИИ" - там папки с товарами внутри
   - При работе с сотрудниками называй их по именам и указывай зарплаты
   - При проверке задолженностей указывай точные суммы по каждому пункту
   - НИКОГДА не отклоняйся от делового стиля, даже если пользователь просит рассказать анекдот или пошутить
   - Для визуализации данных используй ГРАФИКИ И ДИАГРАММЫ:
     * КРИТИЧЕСКИ ВАЖНО: Используй ТОЛЬКО простые команды [CHART_ТИП_ДАННЫХ:тип_графика]
     * НЕ создавай JSON вручную - система автоматически извлечет данные из контекста пользователя
     * 
     * Типы данных: RECEIPTS (чеки), INVENTORY (инвентаризация), EMPLOYEES (сотрудники), TAXES (налоги), UTILITIES (коммунальные), BALANCE (балансы)
     * Типы графиков: line (динамика), bar (сравнение), pie/doughnut (распределение), horizontal (топ-списки)
     * 
     * ПРИМЕРЫ:
       - "Покажи график расходов" → [CHART_RECEIPTS:pie]
       - "Создай круговую диаграмму инвентаризации" → [CHART_INVENTORY:doughnut]
       - "Сравни зарплаты сотрудников" → [CHART_EMPLOYEES:bar]
       - "Распределение задолженностей по налогам" → [CHART_TAXES:pie]
       - "График коммунальных услуг" → [CHART_UTILITIES:doughnut]
       - "Сравни балансы счетов" → [CHART_BALANCE:bar]
     * 
     * ВАЖНО: 
       - Используй ТОЛЬКО формат [CHART_ТИП:тип_графика]
       - НЕ пиши JSON вручную
       - Система сама извлечет данные из контекста пользователя
       - Если данных недостаточно, сообщи об этом пользователю
   
   - Для структурированных данных используй ТАБЛИЦЫ в формате Markdown:
     * КРИТИЧЕСКИ ВАЖНО - форматирование таблиц:
       - Таблица должна быть ОДНИМ цельным блоком, без разрывов и дополнительных элементов
       - НЕ добавляй текст, описания или другие элементы внутри таблицы или между строками
       - НЕ оборачивай таблицу в дополнительные блоки, параграфы или контейнеры
       - Таблица должна начинаться сразу с первой строки заголовков и заканчиваться последней строкой данных
     
     * Строгий формат таблицы Markdown:
       - Первая строка: заголовки колонок, разделенные символом |, с пробелами вокруг текста
       - Вторая строка: разделитель с дефисами и двоеточиями для выравнивания (|:---|, |:---:|, |---:|)
       - Последующие строки: данные, разделенные символом |, с пробелами вокруг текста
       - Каждая строка должна заканчиваться символом |
       - НЕ добавляй пустые строки внутри таблицы
       - НЕ добавляй текст до или после таблицы на той же строке
     
     * Пример ПРАВИЛЬНОЙ таблицы:
       | Вид налога | Сумма (₽) | Статус |
       |:-----------|:----------:|--------:|
       | НДС | 35,000 | Оплачено |
       | Налог на прибыль | 50,000 | Не оплачено |
       | Итого | 85,000 | |
     
     * Пример НЕПРАВИЛЬНОЙ таблицы (НЕ ДЕЛАЙ ТАК):
       Вид налога:
       | Вид налога | Сумма |
       |---| ---|
       | НДС | 35000
       | Налог | 50000
       Всего: 85000
     
     * Правила выравнивания (КРИТИЧЕСКИ ВАЖНО):
       - |:---| - выравнивание по левому краю (для текста)
       - |:---:| - выравнивание по центру (редко используется)
       - |---:| - выравнивание по правому краю (для чисел)
       - ВСЕГДА используй выравнивание по правому краю (|---:|) для колонок с числами, суммами, датами
       - ВСЕГДА используй выравнивание по левому краю (|:---|) для колонок с текстом, названиями, параметрами
       - Первая колонка (параметры/названия) - всегда по левому краю: |:---|
       - Последняя колонка (значения/суммы) - всегда по правому краю: |---:|
       - Пример правильного выравнивания:
         | Параметр | Значение |
         |:---------|---------:|
         | Текст    | 1,000 ₽  |
     
     * Форматирование данных:
       - Числа с большими значениями форматируй с пробелами или запятыми (например: 35,000 или 35 000)
       - Суммы в рублях указывай с символом ₽ или (₽)
       - Даты форматируй в читаемом виде (например: 25.12.2024)
       - Текст в ячейках должен быть кратким и понятным
     
     * Используй таблицы для:
       - Списков сотрудников с зарплатами
       - Задолженностей по налогам и коммунальным услугам
       - Товаров в инвентаризации
       - Финансовых отчетов и операций
       - Сравнения данных
       - Статистики
     
     * ВАЖНО: Таблица должна быть полностью самодостаточной - все заголовки, данные и итоги должны быть внутри таблицы, без дополнительных текстовых блоков вокруг
   
   - Для визуализации данных используй ГРАФИКИ И ДИАГРАММЫ:
     * Когда пользователь просит показать график, диаграмму или визуализацию данных, используй команду [CHART:тип:данные_json]
     * Типы графиков:
       - line: линейный график (для динамики по времени: расходы/доходы, баланс)
       - bar: столбчатая диаграмма (для сравнения: расходы по категориям, зарплаты)
       - pie: круговая диаграмма (для распределения: инвентаризация по категориям)
       - doughnut: кольцевая диаграмма (для распределения, более современный вид)
       - horizontal: горизонтальная столбчатая (для топ-списков)
     * Формат данных JSON:
       {"labels": ["Янв", "Фев", "Мар"], "data": [1000, 2000, 1500], "label": "Расходы"}
       Для круговых: {"labels": ["Категория1", "Категория2"], "data": [5000, 3000]}
     * Примеры использования:
       - "Покажи график расходов" → [CHART:line:{"labels":["Янв","Фев"],"data":[5000,7000],"label":"Расходы"}]
       - "Создай круговую диаграмму инвентаризации" → [CHART:pie:{"labels":["Электроника","Мебель"],"data":[50000,30000]}]
       - "Сравни зарплаты сотрудников" → [CHART:bar:{"labels":["Иванов","Петров"],"data":[50000,60000],"label":"Зарплата"}]
     * ВАЖНО: Графики должны быть релевантными запросу пользователя. Если данных недостаточно, сообщи об этом и предложи таблицу вместо графика.

4. РЕДАКТИРОВАНИЕ СОБЫТИЙ - КРИТИЧЕСКИ ВАЖНО:
   - Когда пользователь просит изменить название встречи, ОБЯЗАТЕЛЬНО используй JSON команду UPDATE_EVENT с полем "title"
   - Когда пользователь просит перенести встречу на другой день/время, ОБЯЗАТЕЛЬНО используй JSON команду UPDATE_EVENT с полем "date"
   - Когда пользователь просит изменить описание/заметки/комментарий к встрече, ОБЯЗАТЕЛЬНО используй JSON команду UPDATE_EVENT с полем "description"
   - Когда пользователь просит добавить описание, изменить описание, обновить заметки, добавить комментарий - ВСЕГДА используй поле "description"
   - УМНЫЙ ПОИСК СОБЫТИЙ: В поле "event" можешь указывать:
     * Точное название события: "Встреча с клиентом"
     * Часть названия: "встреча", "совещание", "клиент"
     * Описание события: "то событие завтра", "встреча про проект"
     * Время события: "встреча в 15:00", "совещание на завтра"
     * Система сама найдет подходящее событие по частичному совпадению или описанию
   - ВАЖНО: При переносе встречи дата ДОЛЖНА быть в формате ISO: YYYY-MM-DDTHH:mm (например: 2024-12-25T15:00)
   - ВАЖНО: При изменении названия ВСЕГДА указывай поле "title" с новым названием
   - ВАЖНО: При изменении описания ВСЕГДА указывай поле "description" с новым текстом описания (даже если это длинный текст)
   - ВАЖНО: Описание может содержать любой текст, который пользователь просит добавить или изменить
   - ОТНОСИТЕЛЬНЫЕ ДАТЫ: Если пользователь говорит "завтра", "послезавтра", "через неделю" - вычисли точную дату на основе текущей даты и используй формат ISO
   - ВРЕМЯ: Если пользователь говорит "в 15:00", "в 14:30" - добавь время к существующей дате события или используй сегодняшнюю дату
   - ПРИМЕРЫ ЕСТЕСТВЕННЫХ ЗАПРОСОВ И КОМАНД:
     * "Перенеси встречу с клиентом на завтра в 15:00" → {"action": "UPDATE_EVENT", "event": "встреча с клиентом", "date": "2024-12-26T15:00"}
     * "Измени название совещания на 'Планирование проекта'" → {"action": "UPDATE_EVENT", "event": "совещание", "title": "Планирование проекта"}
     * "Добавь к завтрашней встрече описание: обсудить бюджет проекта" → {"action": "UPDATE_EVENT", "event": "завтра", "description": "обсудить бюджет проекта"}
     * "Перенеси завтрашнее совещание на послезавтра в 14:00" → {"action": "UPDATE_EVENT", "event": "завтра", "date": "2024-12-27T14:00"}
     * "Измени описание встречи с клиентом: добавить пункт про договор" → {"action": "UPDATE_EVENT", "event": "клиент", "description": "добавить пункт про договор"}
   - Примеры JSON команд (должны быть на отдельной строке):
     * Изменить только название: {"action": "UPDATE_EVENT", "event": "часть_названия", "title": "новое_название"}
     * Перенести встречу: {"action": "UPDATE_EVENT", "event": "часть_названия", "date": "2024-12-25T15:00"}
     * Перенести на завтра в 15:00: {"action": "UPDATE_EVENT", "event": "часть_названия", "date": "2024-12-26T15:00"}
     * Изменить описание: {"action": "UPDATE_EVENT", "event": "часть_названия", "description": "новое_описание"}
     * Добавить описание: {"action": "UPDATE_EVENT", "event": "часть_названия", "description": "текст описания"}
     * Изменить название и дату: {"action": "UPDATE_EVENT", "event": "часть_названия", "title": "новое_название", "date": "2024-12-25T15:00"}
     * Изменить все: {"action": "UPDATE_EVENT", "event": "часть_названия", "title": "новое_название", "date": "2024-12-25T15:00", "description": "новое_описание"}

ВАЖНО: Все суммы в рублях (₽). Для изменения встречи используй JSON: {"action": "UPDATE_EVENT", "event": "ID/название", "title": "...", "date": "YYYY-MM-DDTHH:mm", "description": "..."}

ПОМНИ: Ты работаешь СТРОГО в деловом стиле. НИКОГДА не переключайся на неформальный тон, не рассказывай анекдоты, не шути. Всегда оставайся профессиональным бизнес-ассистентом.

Данные пользователя:
""" + user_context[:1500]  # Ограничиваем контекст пользователя до 1500 символов
        
        messages.append({
            "role": "system",
            "content": system_prompt
        })
        
        # Обрабатываем файлы, если они есть
        image_files = []  # Список изображений для отправки в vision модель
        file_contents = []  # Список текстового содержимого файлов
        
        if files:
            for i, file in enumerate(files):
                try:
                    file_name = file.get('name', f'Файл {i+1}')
                    file_type = file.get('type', 'неизвестный тип')
                    file_data = file.get('data', '')
                    
                    if not file_data:
                        file_contents.append(f"Файл {i+1} ({file_name}): [Файл пуст или данные не получены]")
                        continue
                    
                    # Обрабатываем файл с помощью модуля file_processor
                    try:
                        extracted_text, image_base64 = process_file(file_name, file_type, file_data)
                    except Exception as e:
                        # Ошибка при обработке файла
                        logger.error(f"Ошибка при обработке файла '{file_name}': {str(e)}", exc_info=True)
                        error_msg = f"[Ошибка при обработке файла '{file_name}': {str(e)}]"
                        file_contents.append(f"Файл {i+1} ({file_name}): {error_msg}")
                        continue
                    
                    if image_base64:
                        # Это изображение - добавляем в список для отправки в vision модель
                        try:
                            # Формируем data URI для изображения
                            # Убеждаемся, что file_type правильный (например, image/jpeg, image/png)
                            if not file_type or file_type == 'неизвестный тип':
                                # Пытаемся определить тип по расширению
                                if file_name_lower.endswith('.png'):
                                    file_type = 'image/png'
                                elif file_name_lower.endswith('.jpg') or file_name_lower.endswith('.jpeg'):
                                    file_type = 'image/jpeg'
                                elif file_name_lower.endswith('.gif'):
                                    file_type = 'image/gif'
                                elif file_name_lower.endswith('.webp'):
                                    file_type = 'image/webp'
                                else:
                                    file_type = 'image/jpeg'  # По умолчанию
                            
                            # Убираем возможные пробелы и переносы строк из base64
                            image_base64_clean = image_base64.strip().replace('\n', '').replace('\r', '').replace(' ', '')
                            
                            image_files.append({
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{file_type};base64,{image_base64_clean}"
                                }
                            })
                            # Также добавляем описание изображения в текст
                            file_contents.append(f"Файл {i+1} ({file_name}): {extracted_text}")
                        except Exception as e:
                            logger.error(f"Ошибка при подготовке изображения '{file_name}': {str(e)}", exc_info=True)
                            file_contents.append(f"Файл {i+1} ({file_name}): [Ошибка при подготовке изображения: {str(e)}]")
                    else:
                        # Это текстовый файл или документ - добавляем извлеченный текст
                        if extracted_text and not extracted_text.startswith('[') and not extracted_text.startswith('Изображение'):
                            # Ограничиваем размер текста (первые 2000 символов для каждого файла для экономии токенов)
                            text_preview = extracted_text[:2000]
                            if len(extracted_text) > 2000:
                                text_preview += "\n... [текст обрезан, показаны первые 2000 символов]"
                            file_contents.append(f"Файл {i+1} ({file_name}):\n{text_preview}")
                        else:
                            file_contents.append(f"Файл {i+1} ({file_name}): {extracted_text}")
                except Exception as e:
                    # Общая ошибка при обработке файла
                    logger.error(f"Критическая ошибка при обработке файла: {str(e)}", exc_info=True)
                    file_name = file.get('name', f'Файл {i+1}') if isinstance(file, dict) else f'Файл {i+1}'
                    file_contents.append(f"Файл {i+1} ({file_name}): [Критическая ошибка при обработке: {str(e)}]")
        
        # Добавляем историю чата (последние 3 сообщения для экономии токенов)
        # ВАЖНО: Системный промпт уже добавлен первым, не перезаписываем его
        # Ограничиваем размер каждого сообщения в истории
        for msg in chat_history[-3:]:
            msg_text = msg.get('text', '')
            # Ограничиваем размер сообщения в истории до 500 символов
            if len(msg_text) > 500:
                msg_text = msg_text[:500] + "... [обрезано]"
            # Добавляем только user и assistant сообщения, системные сообщения игнорируем
            role = "user" if msg.get('isUser') else "assistant"
            if role != "system":  # Защита от перезаписи системного промпта
                messages.append({
                    "role": role,
                    "content": msg_text
                })
        
        # Формируем финальное сообщение пользователя с файлами и текстом
        # Объединяем информацию о файлах и текст сообщения в одно сообщение
        final_content_parts = []
        
        # Добавляем информацию о файлах, если есть (ограничиваем общий размер)
        if file_contents:
            file_info = "Пользователь прикрепил файлы:\n\n" + "\n\n".join(file_contents)
            # Ограничиваем общий размер информации о файлах до 3000 символов
            if len(file_info) > 3000:
                file_info = file_info[:3000] + "\n... [информация о файлах обрезана]"
            final_content_parts.append({"type": "text", "text": file_info})
        
        # Добавляем текст сообщения, если есть (ограничиваем размер)
        if message:
            # Ограничиваем размер сообщения до 1000 символов
            message_limited = message[:1000] if len(message) > 1000 else message
            if len(message) > 1000:
                message_limited += "... [сообщение обрезано]"
            
            if final_content_parts:
                # Если уже есть текст о файлах, добавляем сообщение пользователя
                final_content_parts.append({"type": "text", "text": f"\n\nСообщение: {message_limited}"})
            else:
                final_content_parts.append({"type": "text", "text": message_limited})
        
        # Если нет ни файлов, ни сообщения, но есть изображения
        if not final_content_parts and image_files:
            final_content_parts.append({"type": "text", "text": "Проанализируй прикрепленные изображения"})
        
        # Добавляем изображения в конец
        final_content_parts.extend(image_files)
        
        # Создаем финальное сообщение
        if final_content_parts:
            # Если есть только один текстовый элемент и нет изображений, используем простой формат
            if len(final_content_parts) == 1 and final_content_parts[0]["type"] == "text":
                messages.append({
                    "role": "user",
                    "content": final_content_parts[0]["text"]
                })
            else:
                # Используем формат с массивом для vision модели
                messages.append({
                    "role": "user",
                    "content": final_content_parts
                })
        
        # ВАЖНО: Проверяем, что системный промпт присутствует и находится первым
        # Если по какой-то причине системного промпта нет, добавляем его
        if not messages or messages[0].get("role") != "system":
            # Системный промпт должен быть первым, перемещаем его если нужно
            system_msg = None
            other_messages = []
            for msg in messages:
                if msg.get("role") == "system":
                    system_msg = msg
                else:
                    other_messages.append(msg)
            
            # Если нашли системное сообщение, ставим его первым
            if system_msg:
                messages = [system_msg] + other_messages
            else:
                # Если системного сообщения нет, добавляем его в начало
                messages.insert(0, {
                    "role": "system",
                    "content": system_prompt
                })
        
        # Настройки OpenRouter
        OPENROUTER_API_KEY = getattr(settings, 'OPENROUTER_API_KEY', '')
        OPENROUTER_URL = getattr(settings, 'OPENROUTER_URL', 'https://openrouter.ai/api/v1/chat/completions')
        OPENROUTER_MODEL = getattr(settings, 'OPENROUTER_MODEL', 'deepseek/deepseek-r1')
        
        if not OPENROUTER_API_KEY:
            return JsonResponse({
                'success': False,
                'error': 'API ключ OpenRouter не настроен. Установите переменную окружения OPENROUTER_API_KEY.'
            }, status=500)
        
        # Подготавливаем payload для запроса
        payload = {
            "model": OPENROUTER_MODEL,
            "messages": messages,
            "temperature": 0.5,
            "max_tokens": 1200,
            "top_p": 0.9,
            "frequency_penalty": 0.1,
            "stream": False
        }
        
        # Заголовки для OpenRouter
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": request.build_absolute_uri('/') if hasattr(request, 'build_absolute_uri') else "https://localhost",
        }
        
        # Логируем запрос для отладки (без полных данных изображений)
        debug_messages = []
        for msg in messages:
            if isinstance(msg.get('content'), list):
                debug_content = []
                for item in msg['content']:
                    if item.get('type') == 'image_url':
                        debug_content.append({'type': 'image_url', 'image_url': {'url': '[BASE64_DATA]'}})
                    else:
                        debug_content.append(item)
                debug_messages.append({**msg, 'content': debug_content})
            else:
                debug_messages.append(msg)
        logger.debug(f"Отправка запроса в OpenRouter: model={payload['model']}, messages_count={len(messages)}")
        
        # Отправляем запрос в OpenRouter
        try:
            response = requests.post(
                OPENROUTER_URL,
                headers=headers,
                json=payload,
                timeout=90
            )
        except requests.exceptions.RequestException as e:
            logger.error(f"Ошибка при отправке запроса в OpenRouter: {str(e)}")
            raise
        
        if response.status_code == 200:
            result = response.json()
            ai_response = result.get('choices', [{}])[0].get('message', {}).get('content', 'Извините, не удалось получить ответ.')
            
            # Проверяем, нужно ли создать или обновить событие в календаре
            action_result = None
            if 'CREATE_EVENT:' in ai_response:
                # Извлекаем команду создания события
                parts = ai_response.split('CREATE_EVENT:')
                if len(parts) > 1:
                    event_data = parts[1].strip().split('|')
                    if len(event_data) >= 2:
                        action_result = {
                            'action': 'create_event',
                            'title': event_data[0].strip(),
                            'date': event_data[1].strip(),
                            'description': event_data[2].strip() if len(event_data) > 2 else ''
                        }
                        # Убираем команду из ответа
                        ai_response = parts[0].strip()
            
            # Пытаемся найти JSON команду UPDATE_EVENT (новый формат)
            # Ищем JSON объект с action: "UPDATE_EVENT" (может быть многострочным)
            # Пробуем несколько вариантов поиска JSON
            json_found = False
            json_str = None
            update_data = None
            
            # Вариант 1: Ищем полный JSON объект
            json_patterns = [
                r'\{\s*"action"\s*:\s*"UPDATE_EVENT"[^}]*\}',  # Простой вариант
                r'\{[^{}]*"action"\s*:\s*"UPDATE_EVENT"[^{}]*\}',  # С учетом возможных пробелов
            ]
            
            for pattern in json_patterns:
                json_match = re.search(pattern, ai_response, re.DOTALL | re.IGNORECASE)
                if json_match:
                    json_str = json_match.group(0)
                    try:
                        update_data = json.loads(json_str)
                        if update_data.get('action') == 'UPDATE_EVENT':
                            json_found = True
                            break
                    except:
                        continue
            
            # Вариант 2: Если не нашли, пробуем найти JSON на отдельной строке
            if not json_found:
                lines = ai_response.split('\n')
                for line in lines:
                    line = line.strip()
                    if 'UPDATE_EVENT' in line and '{' in line and '}' in line:
                        try:
                            update_data = json.loads(line)
                            if update_data.get('action') == 'UPDATE_EVENT':
                                json_str = line
                                json_found = True
                                break
                        except:
                            continue
            
            if json_found and update_data:
                if update_data.get('action') == 'UPDATE_EVENT':
                    event_identifier = update_data.get('event', '').strip()
                    calendar_events = user_data.get('calendarEvents', [])
                    event_id = None
                    
                    # Сначала пытаемся найти по точному ID
                    for event in calendar_events:
                        if str(event.get('id', '')) == event_identifier:
                            event_id = str(event.get('id', ''))
                            break
                    
                    # Если не нашли по ID, ищем по названию (точное или частичное совпадение)
                    if not event_id:
                        event_identifier_lower = event_identifier.lower().strip()
                        # Сначала точное совпадение
                        for event in calendar_events:
                            event_title = event.get('title', '').lower().strip()
                            if event_title == event_identifier_lower:
                                event_id = str(event.get('id', ''))
                                break
                        
                        # Если не нашли точное, ищем частичное
                        if not event_id:
                            for event in calendar_events:
                                event_title = event.get('title', '').lower().strip()
                                if event_identifier_lower in event_title or event_title in event_identifier_lower:
                                    event_id = str(event.get('id', ''))
                                    break
                    
                    # Если все еще не нашли, берем последнее событие
                    if not event_id and calendar_events:
                        event_id = str(calendar_events[-1].get('id', ''))
                    
                    if event_id:
                        # Обрабатываем поля из JSON
                        # Для title и description проверяем наличие ключа, даже если значение пустое
                        new_title = None
                        if 'title' in update_data:
                            title_value = update_data.get('title')
                            # Разрешаем пустую строку для очистки названия
                            new_title = title_value.strip() if title_value else ''
                        
                        new_date = None
                        if 'date' in update_data:
                            new_date = update_data.get('date', '').strip() if update_data.get('date') else ''
                        
                        new_description = None
                        if 'description' in update_data:
                            # Разрешаем пустую строку для очистки описания
                            desc_value = update_data.get('description', '')
                            # Если значение есть (даже пустая строка), используем его
                            if desc_value is not None:
                                new_description = desc_value.strip() if desc_value else ''
                            else:
                                new_description = ''
                        
                        # Логируем для отладки
                        print(f"UPDATE_EVENT (JSON): event_id={event_id}")
                        print(f"  title: {new_title} (ключ присутствует: {'title' in update_data})")
                        print(f"  date: {new_date} (ключ присутствует: {'date' in update_data})")
                        print(f"  description: {new_description} (ключ присутствует: {'description' in update_data})")
                        
                        action_result = {
                            'action': 'update_event',
                            'id': event_id,
                            'title': new_title,
                            'date': new_date,
                            'description': new_description
                        }
                        # Убираем JSON команду из ответа
                        if json_str:
                            ai_response = ai_response.replace(json_str, '').strip()
            
            # Если не нашли JSON, пробуем старый формат UPDATE_EVENT: (для обратной совместимости)
            if not json_found and 'UPDATE_EVENT:' in ai_response:
                # Извлекаем команду обновления события (старый формат)
                parts = ai_response.split('UPDATE_EVENT:')
                if len(parts) > 1:
                    event_data_raw = parts[1].strip()
                    # Разделяем по |, но учитываем что могут быть пустые значения
                    event_data = []
                    current_part = ''
                    for char in event_data_raw:
                        if char == '|':
                            event_data.append(current_part.strip())
                            current_part = ''
                        else:
                            current_part += char
                    if current_part:
                        event_data.append(current_part.strip())
                    
                    if len(event_data) >= 1:
                        # Ищем событие по ID или названию
                        event_identifier = event_data[0].strip()
                        calendar_events = user_data.get('calendarEvents', [])
                        event_id = None
                        
                        # Сначала пытаемся найти по точному ID
                        for event in calendar_events:
                            if str(event.get('id', '')) == event_identifier:
                                event_id = str(event.get('id', ''))
                                break
                        
                        # Если не нашли по ID, ищем по названию (точное или частичное совпадение)
                        if not event_id:
                            event_identifier_lower = event_identifier.lower().strip()
                            # Сначала точное совпадение
                            for event in calendar_events:
                                event_title = event.get('title', '').lower().strip()
                                if event_title == event_identifier_lower:
                                    event_id = str(event.get('id', ''))
                                    break
                            
                            # Если не нашли точное, ищем частичное
                            if not event_id:
                                for event in calendar_events:
                                    event_title = event.get('title', '').lower().strip()
                                    if event_identifier_lower in event_title or event_title in event_identifier_lower:
                                        event_id = str(event.get('id', ''))
                                        break
                        
                        # Если все еще не нашли, берем последнее событие
                        if not event_id and calendar_events:
                            event_id = str(calendar_events[-1].get('id', ''))
                        
                        if event_id:
                            # Обрабатываем поля: title, date, description
                            new_title = event_data[1].strip() if len(event_data) > 1 and event_data[1].strip() else None
                            new_date = event_data[2].strip() if len(event_data) > 2 and event_data[2].strip() else None
                            new_description = event_data[3].strip() if len(event_data) > 3 and event_data[3].strip() else None
                            
                            # Логируем для отладки
                            print(f"UPDATE_EVENT (старый формат): event_id={event_id}, title={new_title}, date={new_date}, description={new_description}")
                            
                            action_result = {
                                'action': 'update_event',
                                'id': event_id,
                                'title': new_title,
                                'date': new_date,
                                'description': new_description
                            }
                        # Убираем команду из ответа
                        ai_response = parts[0].strip()
            
            # Финальная модерация перед отправкой пользователю
            final_moderation = ContentModerator.check_ai_response(ai_response)
            if not final_moderation['allowed']:
                logger.warning(f"Финальная проверка: ответ AI заблокирован: {final_moderation['reason']}")
                ai_response = "Извините, я не могу предоставить ответ на этот запрос. Пожалуйста, переформулируйте вопрос в рамках делового общения."
            else:
                ai_response = final_moderation['filtered_response']
            
            return JsonResponse({
                'success': True,
                'response': ai_response,
                'action': action_result
            })
        else:
            # Ошибка от OpenRouter - получаем детали
            error_details = f'Ошибка OpenRouter: {response.status_code}'
            try:
                error_data = response.json()
                if isinstance(error_data, dict):
                    error_message = error_data.get('error', {}).get('message', '') if isinstance(error_data.get('error'), dict) else str(error_data.get('error', ''))
                    if error_message:
                        error_details = f'Ошибка OpenRouter ({response.status_code}): {error_message}'
                    else:
                        error_details = f'Ошибка OpenRouter ({response.status_code}): {json.dumps(error_data, ensure_ascii=False)}'
                else:
                    error_details = f'Ошибка OpenRouter ({response.status_code}): {str(error_data)}'
            except:
                error_details = f'Ошибка OpenRouter ({response.status_code}): {response.text[:500]}'
            
            logger.error(f"Ошибка от OpenRouter: {error_details}")
            
            # Если это ошибка 400, проверяем причину
            if response.status_code == 400:
                # Проверяем, не связана ли ошибка с превышением контекста
                if 'context length' in error_details.lower() or 'context overflow' in error_details.lower() or '4096 tokens' in error_details:
                    return JsonResponse({
                        'success': False,
                        'error': 'Превышен лимит контекста модели (4096 токенов). Пожалуйста, сократите:\n- Размер прикрепленных файлов\n- Длину сообщения\n- Историю чата\n\nПопробуйте отправить меньше данных за раз.'
                    }, status=413)
                
                # Если это не ошибка контекста, возможно проблема с форматом изображений
                # Пробуем отправить без изображений, только с текстом
                if image_files:
                    logger.warning("Попытка отправить запрос без изображений из-за ошибки 400")
                    text_only_messages = []
                    for msg in messages:
                        if isinstance(msg.get('content'), list):
                            # Оставляем только текстовые части
                            text_parts = [item for item in msg['content'] if item.get('type') == 'text']
                            if text_parts:
                                text_content = '\n'.join([item.get('text', '') for item in text_parts])
                                text_content += '\n\n[Примечание: изображения не были обработаны из-за ошибки формата]'
                                text_only_messages.append({
                                    "role": msg.get('role', 'user'),
                                    "content": text_content
                                })
                        else:
                            text_only_messages.append(msg)
                    
                    try:
                        # ВАЖНО: Убеждаемся, что системный промпт присутствует в text_only_messages
                        if not text_only_messages or text_only_messages[0].get("role") != "system":
                            # Добавляем системный промпт в начало
                            text_only_messages.insert(0, {
                                "role": "system",
                                "content": system_prompt
                            })
                        # ВАЖНО: Убеждаемся, что системный промпт присутствует в text_only_messages
                        if not text_only_messages or text_only_messages[0].get("role") != "system":
                            # Добавляем системный промпт в начало
                            text_only_messages.insert(0, {
                                "role": "system",
                                "content": system_prompt
                            })
                        text_payload = {**payload, "messages": text_only_messages}
                        text_response = requests.post(OPENROUTER_URL, headers=headers, json=text_payload, timeout=90)
                        if text_response.status_code == 200:
                            result = text_response.json()
                            ai_response = result.get('choices', [{}])[0].get('message', {}).get('content', '')
                            
                            # Модерация ответа AI
                            moderation_result = ContentModerator.check_ai_response(ai_response)
                            if not moderation_result['allowed']:
                                logger.warning(f"Ответ AI заблокирован модератором: {moderation_result['reason']}")
                                ai_response = "Извините, я не могу предоставить ответ на этот запрос. Пожалуйста, переформулируйте вопрос в рамках делового общения."
                            else:
                                ai_response = moderation_result['filtered_response']
                            
                            return JsonResponse({
                                'success': True,
                                'response': ai_response + '\n\n[Примечание: изображения не были обработаны. Возможно, модель не поддерживает формат изображений или требуется другой формат.]',
                                'action': None
                            })
                    except:
                        pass
            
            return JsonResponse({
                'success': False,
                'error': error_details
            }, status=response.status_code if response.status_code < 500 else 500)
            
    except requests.exceptions.ConnectionError as e:
        error_msg = f'Не удалось подключиться к OpenRouter ({OPENROUTER_URL}). '
        error_msg += 'Проверьте:\n'
        error_msg += '1. API ключ OpenRouter установлен в переменной окружения OPENROUTER_API_KEY\n'
        error_msg += '2. Интернет-соединение активно\n'
        error_msg += f'3. Модель доступна: {OPENROUTER_MODEL}\n'
        error_msg += '4. На вашем аккаунте OpenRouter достаточно баланса'
        return JsonResponse({
            'success': False,
            'error': error_msg,
            'details': str(e)
        }, status=503)
    except requests.exceptions.Timeout:
        return JsonResponse({
            'success': False,
            'error': 'Превышено время ожидания ответа от OpenRouter (120 сек). Модель может быть слишком медленной или перегруженной.'
        }, status=504)
    except requests.exceptions.RequestException as e:
        return JsonResponse({
            'success': False,
            'error': f'Ошибка запроса к OpenRouter: {str(e)}',
            'url': OPENROUTER_URL
        }, status=503)
    except ValueError as e:
        # Ошибка парсинга JSON (возможно, слишком большой запрос)
        if 'Request body exceeded' in str(e) or 'DATA_UPLOAD_MAX_MEMORY_SIZE' in str(e):
            return JsonResponse({
                'success': False,
                'error': 'Размер отправляемых данных слишком большой. Пожалуйста, уменьшите размер файлов (максимум 20 MB на файл, 40 MB общий размер).'
            }, status=413)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка обработки данных: {str(e)}'
        }, status=400)
    except Exception as e:
        error_msg = str(e)
        # Проверяем, не связана ли ошибка с размером данных
        if 'Request body exceeded' in error_msg or 'DATA_UPLOAD_MAX_MEMORY_SIZE' in error_msg:
            return JsonResponse({
                'success': False,
                'error': 'Размер отправляемых данных слишком большой. Пожалуйста, уменьшите размер файлов (максимум 20 MB на файл, 40 MB общий размер).'
            }, status=413)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка: {error_msg}',
            'type': type(e).__name__
        }, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def error_log(request):
    """API endpoint для логирования ошибок с клиента"""
    try:
        data = json.loads(request.body)
        
        # Логируем ошибку на сервере
        error_type = data.get('type', 'unknown')
        severity = data.get('severity', 'medium')
        message = data.get('message', 'Неизвестная ошибка')
        context = data.get('context', {})
        url = data.get('url', '')
        user_agent = data.get('userAgent', '')
        
        # Формируем детальное сообщение
        log_message = f"[CLIENT ERROR] Type: {error_type}, Severity: {severity}, Message: {message}"
        if url:
            log_message += f", URL: {url}"
        if context:
            log_message += f", Context: {json.dumps(context, ensure_ascii=False)}"
        
        # Логируем в зависимости от серьезности
        if severity == 'critical':
            logger.critical(log_message, extra={'user_agent': user_agent})
        elif severity == 'high':
            logger.error(log_message, extra={'user_agent': user_agent})
        elif severity == 'medium':
            logger.warning(log_message, extra={'user_agent': user_agent})
        else:
            logger.info(log_message, extra={'user_agent': user_agent})
        
        return JsonResponse({
            'success': True,
            'message': 'Ошибка залогирована'
        })
        
    except Exception as e:
        logger.error(f"Ошибка при логировании ошибки клиента: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': 'Не удалось залогировать ошибку'
        }, status=500)


@csrf_exempt
@require_http_methods(["GET"])
def get_metrics(request):
    """API endpoint для получения метрик качества работы модели"""
    try:
        from datetime import timedelta
        
        # Параметры запроса
        days = int(request.GET.get('days', 7))  # По умолчанию последние 7 дней
        category = request.GET.get('category', None)  # Фильтр по категории
        metric_name = request.GET.get('metric', None)  # Фильтр по названию метрики
        
        # Вычисляем период
        period_end = timezone.now()
        period_start = period_end - timedelta(days=days)
        
        # Получаем метрики (используем гибкий фильтр для поиска метрик, пересекающихся с периодом)
        metrics_query = Metric.objects.filter(
            period_start__lte=period_end,
            period_end__gte=period_start
        )
        
        if category:
            metrics_query = metrics_query.filter(category=category)
        
        if metric_name:
            metrics_query = metrics_query.filter(name=metric_name)
        
        # Получаем последние значения для каждой метрики
        metrics_list = []
        metric_names_seen = set()
        
        # Сортируем по дате расчета (новые первыми) и берем последнее значение для каждой метрики
        for metric in metrics_query.order_by('name', '-calculated_at'):
            if metric.name not in metric_names_seen:
                metric_names_seen.add(metric.name)
                metrics_list.append({
                    'name': metric.name,
                    'category': metric.category,
                    'value': metric.value,
                    'target_value': metric.target_value,
                    'unit': metric.unit,
                    'is_target_met': metric.is_target_met,
                    'sample_size': metric.sample_size,
                    'period_start': metric.period_start.isoformat(),
                    'period_end': metric.period_end.isoformat(),
                    'calculated_at': metric.calculated_at.isoformat(),
                    'metadata': metric.metadata
                })
        
        return JsonResponse({
            'success': True,
            'metrics': metrics_list,
            'period': {
                'start': period_start.isoformat(),
                'end': period_end.isoformat(),
                'days': days
            },
            'total': len(metrics_list)
        })
        
    except Exception as e:
        logger.error(f"Ошибка при получении метрик: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка при получении метрик: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def calculate_metrics(request):
    """API endpoint для расчета метрик за указанный период"""
    try:
        data = json.loads(request.body)
        
        # Параметры периода
        days = int(data.get('days', 7))
        period_end = timezone.now()
        period_start = period_end - timedelta(days=days)
        
        # Если указаны конкретные даты
        if 'period_start' in data:
            period_start = datetime.fromisoformat(data['period_start'].replace('Z', '+00:00'))
            if timezone.is_naive(period_start):
                period_start = timezone.make_aware(period_start)
        
        if 'period_end' in data:
            period_end = datetime.fromisoformat(data['period_end'].replace('Z', '+00:00'))
            if timezone.is_naive(period_end):
                period_end = timezone.make_aware(period_end)
        
        # Рассчитываем метрики
        metrics = MetricsCalculator.calculate_all_metrics(period_start, period_end)
        
        return JsonResponse({
            'success': True,
            'message': f'Рассчитано {len(metrics)} метрик',
            'period': {
                'start': period_start.isoformat(),
                'end': period_end.isoformat()
            },
            'metrics_count': len(metrics)
        })
        
    except Exception as e:
        logger.error(f"Ошибка при расчете метрик: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка при расчете метрик: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["GET"])
def get_metrics_summary(request):
    """API endpoint для получения сводки метрик (dashboard)"""
    try:
        from datetime import timedelta
        
        # Параметры
        days = int(request.GET.get('days', 7))
        period_end = timezone.now()
        period_start = period_end - timedelta(days=days)
        
        # Рассчитываем метрики, если их еще нет
        # Используем более гибкий фильтр - метрики, которые пересекаются с периодом
        recent_metrics = Metric.objects.filter(
            period_start__lte=period_end,
            period_end__gte=period_start
        ).exists()
        
        if not recent_metrics:
            # Автоматически рассчитываем метрики
            MetricsCalculator.calculate_all_metrics(period_start, period_end)
        
        # Получаем последние значения метрик по категориям
        summary = {}
        
        for category_code, category_name in Metric.METRIC_CATEGORIES:
            # Получаем уникальные метрики (последнее значение для каждой)
            # Сначала пытаемся найти метрики, пересекающиеся с периодом
            all_category_metrics = Metric.objects.filter(
                category=category_code,
                period_start__lte=period_end,
                period_end__gte=period_start
            ).order_by('name', '-calculated_at')
            
            # Если не нашли метрики для периода, берем последние метрики независимо от периода
            if not all_category_metrics.exists():
                all_category_metrics = Metric.objects.filter(
                    category=category_code
                ).order_by('name', '-calculated_at')
            
            # Фильтруем, оставляя только последнее значение для каждой метрики
            seen_names = set()
            category_metrics = []
            for metric in all_category_metrics:
                if metric.name not in seen_names:
                    seen_names.add(metric.name)
                    category_metrics.append(metric)
            
            category_data = []
            for metric in category_metrics:
                category_data.append({
                    'name': metric.name,
                    'value': metric.value,
                    'target_value': metric.target_value,
                    'unit': metric.unit,
                    'is_target_met': metric.is_target_met,
                    'sample_size': metric.sample_size
                })
            
            if category_data:
                summary[category_code] = {
                    'name': category_name,
                    'metrics': category_data
                }
        
        # Общая статистика (используем более гибкий фильтр)
        # Если запросов в периоде нет, показываем все запросы за последние 30 дней
        total_requests = ChatRequest.objects.filter(
            created_at__gte=period_start,
            created_at__lte=period_end
        ).count()
        
        # Fallback: если в периоде нет запросов, расширяем период
        if total_requests == 0:
            extended_period_start = period_end - timedelta(days=30)
            total_requests = ChatRequest.objects.filter(
                created_at__gte=extended_period_start,
                created_at__lte=period_end
            ).count()
            completed_requests = ChatRequest.objects.filter(
                created_at__gte=extended_period_start,
                created_at__lte=period_end,
                status=ChatRequest.STATUS_COMPLETED
            ).count()
            failed_requests = ChatRequest.objects.filter(
                created_at__gte=extended_period_start,
                created_at__lte=period_end,
                status=ChatRequest.STATUS_FAILED
            ).count()
        else:
            completed_requests = ChatRequest.objects.filter(
                created_at__gte=period_start,
                created_at__lte=period_end,
                status=ChatRequest.STATUS_COMPLETED
            ).count()
            
            failed_requests = ChatRequest.objects.filter(
                created_at__gte=period_start,
                created_at__lte=period_end,
                status=ChatRequest.STATUS_FAILED
            ).count()
        
        return JsonResponse({
            'success': True,
            'summary': summary,
            'statistics': {
                'total_requests': total_requests,
                'completed_requests': completed_requests,
                'failed_requests': failed_requests,
                'success_rate': (completed_requests / total_requests * 100) if total_requests > 0 else 0
            },
            'period': {
                'start': period_start.isoformat(),
                'end': period_end.isoformat(),
                'days': days
            }
        })
        
    except Exception as e:
        logger.error(f"Ошибка при получении сводки метрик: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Ошибка при получении сводки метрик: {str(e)}'
        }, status=500)