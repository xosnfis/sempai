"""
Модуль для обработки и извлечения текста из различных форматов файлов
"""
import base64
import io
import os
from typing import Dict, Optional, Tuple


def extract_text_from_pdf(file_data: str) -> str:
    """Извлекает текст из PDF файла"""
    try:
        import PyPDF2
        # Убираем префикс data:type;base64, если есть
        if ',' in file_data:
            file_data = file_data.split(',')[1]
        
        # Декодируем base64
        pdf_bytes = base64.b64decode(file_data)
        pdf_file = io.BytesIO(pdf_bytes)
        
        # Читаем PDF
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_parts = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(f"--- Страница {page_num + 1} ---\n{text}\n")
            except Exception as e:
                text_parts.append(f"--- Страница {page_num + 1} ---\n[Ошибка чтения страницы: {str(e)}]\n")
        
        return "\n".join(text_parts) if text_parts else "[Не удалось извлечь текст из PDF]"
    except ImportError:
        try:
            import pdfplumber
            # Убираем префикс data:type;base64, если есть
            if ',' in file_data:
                file_data = file_data.split(',')[1]
            
            # Декодируем base64
            pdf_bytes = base64.b64decode(file_data)
            pdf_file = io.BytesIO(pdf_bytes)
            
            text_parts = []
            with pdfplumber.open(pdf_file) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            text_parts.append(f"--- Страница {page_num + 1} ---\n{text}\n")
                    except Exception as e:
                        text_parts.append(f"--- Страница {page_num + 1} ---\n[Ошибка чтения страницы: {str(e)}]\n")
            
            return "\n".join(text_parts) if text_parts else "[Не удалось извлечь текст из PDF]"
        except ImportError:
            return "[Библиотека для чтения PDF не установлена. Установите PyPDF2 или pdfplumber]"
    except Exception as e:
        return f"[Ошибка при чтении PDF: {str(e)}]"


def extract_text_from_docx(file_data: str) -> str:
    """Извлекает текст из DOCX файла"""
    try:
        from docx import Document
        
        # Убираем префикс data:type;base64, если есть
        if ',' in file_data:
            file_data = file_data.split(',')[1]
        
        # Декодируем base64
        docx_bytes = base64.b64decode(file_data)
        docx_file = io.BytesIO(docx_bytes)
        
        # Читаем DOCX
        doc = Document(docx_file)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Также извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts) if text_parts else "[Документ пуст или не удалось извлечь текст]"
    except ImportError:
        return "[Библиотека python-docx не установлена]"
    except Exception as e:
        return f"[Ошибка при чтении DOCX: {str(e)}]"


def extract_text_from_xlsx(file_data: str) -> str:
    """Извлекает текст из XLSX файла"""
    try:
        from openpyxl import load_workbook
        
        # Убираем префикс data:type;base64, если есть
        if ',' in file_data:
            file_data = file_data.split(',')[1]
        
        # Декодируем base64
        xlsx_bytes = base64.b64decode(file_data)
        xlsx_file = io.BytesIO(xlsx_bytes)
        
        # Читаем XLSX
        workbook = load_workbook(xlsx_file, data_only=True)
        text_parts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"--- Лист: {sheet_name} ---")
            
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                row_text = " | ".join(row_data)
                if row_text.strip() and not all(not cell.strip() for cell in row_data):
                    text_parts.append(row_text)
            
            text_parts.append("")  # Пустая строка между листами
        
        return "\n".join(text_parts) if text_parts else "[Таблица пуста или не удалось извлечь данные]"
    except ImportError:
        return "[Библиотека openpyxl не установлена]"
    except Exception as e:
        return f"[Ошибка при чтении XLSX: {str(e)}]"


def extract_text_from_text_file(file_data: str) -> str:
    """Извлекает текст из текстового файла"""
    try:
        # Убираем префикс data:type;base64, если есть
        if ',' in file_data:
            file_data = file_data.split(',')[1]
        
        # Декодируем base64
        decoded = base64.b64decode(file_data).decode('utf-8', errors='ignore')
        return decoded
    except Exception as e:
        return f"[Ошибка при чтении текстового файла: {str(e)}]"


def process_file(file_name: str, file_type: str, file_data: str) -> Tuple[str, Optional[str]]:
    """
    Обрабатывает файл и извлекает из него текст или возвращает данные для изображения
    
    Returns:
        Tuple[str, Optional[str]]: (extracted_text, image_base64)
        - extracted_text: извлеченный текст или описание файла
        - image_base64: base64 данные изображения (если это изображение) или None
    """
    # Проверка входных данных
    if not file_data:
        return "[Файл пуст или данные не получены]", None
    
    if not file_name:
        file_name = "Неизвестный файл"
    
    if not file_type:
        file_type = "application/octet-stream"
    
    file_name_lower = file_name.lower()
    
    # Обработка изображений
    if file_type.startswith('image/'):
        # Для изображений возвращаем base64 данные
        if ',' in file_data:
            # Убираем префикс data:image/...;base64,
            image_base64 = file_data.split(',')[1]
        else:
            image_base64 = file_data
        
        # Пытаемся получить информацию об изображении
        try:
            from PIL import Image
            import io as image_io
            
            # Декодируем для проверки
            image_bytes = base64.b64decode(image_base64)
            img = Image.open(image_io.BytesIO(image_bytes))
            
            # Определяем формат изображения
            img_format = img.format or 'unknown'
            image_info = f"Изображение: формат {img_format}, размер {img.width}x{img.height} пикселей"
            return image_info, image_base64
        except ImportError:
            # Pillow не установлен, но это не критично
            return "Изображение (библиотека для анализа не установлена)", image_base64
        except Exception as e:
            # Ошибка при анализе, но base64 данные валидны
            return f"Изображение (не удалось проанализировать: {str(e)})", image_base64
    
    # Обработка PDF
    if file_name_lower.endswith('.pdf') or file_type == 'application/pdf':
        text = extract_text_from_pdf(file_data)
        return text, None
    
    # Обработка DOCX
    if (file_name_lower.endswith('.docx') or 
        file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
        text = extract_text_from_docx(file_data)
        return text, None
    
    # Обработка DOC (старый формат Word)
    if file_name_lower.endswith('.doc') or file_type == 'application/msword':
        return "[Формат .doc не поддерживается напрямую. Пожалуйста, конвертируйте файл в .docx или .pdf]", None
    
    # Обработка XLSX
    if (file_name_lower.endswith('.xlsx') or 
        file_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'):
        text = extract_text_from_xlsx(file_data)
        return text, None
    
    # Обработка XLS (старый формат Excel)
    if file_name_lower.endswith('.xls') or file_type == 'application/vnd.ms-excel':
        return "[Формат .xls не поддерживается напрямую. Пожалуйста, конвертируйте файл в .xlsx или .csv]", None
    
    # Обработка текстовых файлов
    if (file_type.startswith('text/') or 
        any(file_name_lower.endswith(ext) for ext in ['.txt', '.csv', '.json', '.xml', '.html', '.md', '.py', '.js', '.css', '.log', '.yaml', '.yml'])):
        text = extract_text_from_text_file(file_data)
        return text, None
    
    # Для неизвестных форматов пытаемся прочитать как текст
    try:
        text = extract_text_from_text_file(file_data)
        if text and len(text) > 50:  # Если получили достаточно текста
            return text, None
    except:
        pass
    
    return f"[Формат файла '{file_type}' не поддерживается для извлечения текста]", None

